# ============================================================
# OLIVETTI CREATIVE EDITING PARTNER
# Professional-Grade AI Author Engine
# ============================================================
#
# SYSTEM ARCHITECTURE (Spiderweb of Intelligent, Trainable Functions):
#
# ┌─────────────────────────────────────────────────────────────────┐
# │                    USER TRAINING INPUTS                          │
# ├─────────────────────────────────────────────────────────────────┤
# │ 1. Style Banks → Upload/paste samples → Per-lane training       │
# │ 2. Voice Vault → Add voice samples → Semantic vector storage    │
# │ 3. Story Bible → Import/generate → Canon enforcement            │
# │ 4. Voice Bible → Enable controls → Real-time modulation         │
# └─────────────────────────────────────────────────────────────────┘
#                              ↓
# ┌─────────────────────────────────────────────────────────────────┐
# │                   INTELLIGENT RETRIEVAL                          │
# ├─────────────────────────────────────────────────────────────────┤
# │ • retrieve_style_exemplars() ← Context-aware semantic search    │
# │ • retrieve_mixed_exemplars() ← Voice vault vector matching      │
# │ • _story_bible_text() ← Canon assembly                          │
# │ • engine_style_directive() ← Style guidance generation          │
# └─────────────────────────────────────────────────────────────────┘
#                              ↓
# ┌─────────────────────────────────────────────────────────────────┐
# │              UNIFIED AI BRIEF CONSTRUCTION                       │
# ├─────────────────────────────────────────────────────────────────┤
# │ build_partner_brief(action, lane):                              │
# │   ✓ AI Intensity → temperature_from_intensity()                 │
# │   ✓ Writing Style → engine_style_directive() + exemplars        │
# │   ✓ Genre Influence → mood/pacing directives                    │
# │   ✓ Trained Voice → semantic retrieval of user samples          │
# │   ✓ Match My Style → one-shot adaptation                        │
# │   ✓ Voice Lock → hard constraints (MANDATORY)                   │
# │   ✓ POV/Tense → technical specs                                 │
# │   ✓ Story Bible → canon enforcement                             │
# │   ✓ Lane Detection → context-aware mode (Dialogue/Narration)    │
# └─────────────────────────────────────────────────────────────────┘
#                              ↓
# ┌─────────────────────────────────────────────────────────────────┐
# │                 UNIFIED AI GENERATION                            │
# ├─────────────────────────────────────────────────────────────────┤
# │ call_openai(brief, task, text):                                 │
# │   → OpenAI API with temperature from AI Intensity               │
# │   → System prompt = full Voice Bible brief                      │
# │   → User prompt = action-specific task                          │
# │   → Returns professional-grade prose                            │
# └─────────────────────────────────────────────────────────────────┘
#                              ↓
# ┌─────────────────────────────────────────────────────────────────┐
# │               ALL ACTIONS USE SAME ENGINE                        │
# ├─────────────────────────────────────────────────────────────────┤
# │ • partner_action(Write/Rewrite/Expand/etc) ← Writing desk       │
# │ • generate_story_bible_section() ← Story Bible generation       │
# │ • sb_breakdown_ai() ← Import document parsing                   │
# │ ALL route through build_partner_brief() + call_openai()         │
# └─────────────────────────────────────────────────────────────────┘
#                              ↓
# ┌─────────────────────────────────────────────────────────────────┐
# │                  ADAPTIVE FEEDBACK LOOP                          │
# ├─────────────────────────────────────────────────────────────────┤
# │ User adds training → Vector storage → Semantic retrieval →      │
# │ → Better exemplars → Improved AI brief → Higher quality prose   │
# │                                                                  │
# │ CONTINUOUS IMPROVEMENT: More training = Better adaptation        │
# └─────────────────────────────────────────────────────────────────┘
#
# VERIFICATION CHECKLIST (All ✓ = Fully Connected):
# ✓ AI Intensity → temperature_from_intensity() → call_openai()
# ✓ Style Banks → retrieve_style_exemplars() → build_partner_brief()
# ✓ Voice Vault → retrieve_mixed_exemplars() → build_partner_brief()
# ✓ Story Bible → _story_bible_text() → build_partner_brief()
# ✓ Voice Bible Controls → build_partner_brief() → ALL AI actions
# ✓ Write/Rewrite/Expand → partner_action() → build_partner_brief()
# ✓ Generate Synopsis/etc → generate_story_bible_section() → Voice Bible
# ✓ Import → sb_breakdown_ai() → call_openai() → Voice Bible
# ✓ Status Display → _get_voice_bible_summary() → Real-time feedback
# ✓ Bay System → Project management → Story Bible linking
#
# DATA FLOW EXAMPLE (Write Action):
# User clicks "Write" 
#   → queue_action("Write")
#   → partner_action("Write")
#   → lane = current_lane_from_draft(text)  [Context detection]
#   → brief = build_partner_brief("Write", lane)
#        → ai_x = st.session_state.ai_intensity  [0.75]
#        → story_bible = _story_bible_text()  [Canon]
#        → style_exemplars = retrieve_style_exemplars(...)  [Semantic search]
#        → voice_exemplars = retrieve_mixed_exemplars(...)  [Vector match]
#        → Returns: 500-line unified prompt with ALL controls
#   → out = call_openai(brief, task, text)
#        → temperature = temperature_from_intensity(0.75) = 0.86
#        → OpenAI API call with full brief
#        → Returns: Professional prose matching ALL Voice Bible settings
#   → apply_append(out)  [Add to draft]
#   → status = "Write complete (AI:HIGH • Style:LYRICAL • Genre:Literary)"
#
# RESULT: Spiderweb fully connected. Every component feeds into unified AI.
# ============================================================


import os
import re
import math
import json
import hashlib
import logging
from io import BytesIO
from datetime import datetime
from typing import List, Tuple, Dict, Any, Optional

import streamlit as st

# ============================================================
# OLIVETTI DESK — one file, production-stable, paste+click
# ============================================================

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("Olivetti")


# ============================================================
# ENV / METADATA HYGIENE
# ============================================================
os.environ.setdefault("MS_APP_ID", "olivetti-writing-desk")
os.environ.setdefault("ms-appid", "olivetti-writing-desk")

DEFAULT_MODEL = "gpt-4o-mini"

def _get_openai_key_or_empty() -> str:
    try:
        return str(st.secrets.get("OPENAI_API_KEY", ""))  # type: ignore[attr-defined]
    except Exception:
        return ""

def _get_openai_model() -> str:
    try:
        return str(st.secrets.get("OPENAI_MODEL", DEFAULT_MODEL))  # type: ignore[attr-defined]
    except Exception:
        return os.getenv("OPENAI_MODEL", DEFAULT_MODEL)

OPENAI_MODEL = _get_openai_model()

def has_openai_key() -> bool:
    return bool(os.getenv("OPENAI_API_KEY") or _get_openai_key_or_empty())

def require_openai_key() -> str:
    """Stop the app with a clear message if no OpenAI key is configured."""
    key = os.getenv("OPENAI_API_KEY") or _get_openai_key_or_empty()
    if not key:
        st.error(
            "OPENAI_API_KEY is not set. Add it as an environment variable (OPENAI_API_KEY) or as a Streamlit secret."
        )
        st.stop()
    return key




# ============================================================
# GLOBALS
# ============================================================
LANES = ["Dialogue", "Narration", "Interiority", "Action"]
BAYS = ["NEW", "ROUGH", "EDIT", "FINAL"]


ENGINE_STYLES = ["NARRATIVE", "DESCRIPTIVE", "EMOTIONAL", "LYRICAL"]
AUTOSAVE_DIR = "autosave"
AUTOSAVE_PATH = os.path.join(AUTOSAVE_DIR, "olivetti_state.json")
MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10MB guardrail

WORD_RE = re.compile(r"[A-Za-z']+")


# ============================================================
# UTILS
# ============================================================
def now_ts() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _normalize_text(s: str) -> str:
    t = (s or "").strip()
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    t = re.sub(r"\n{3,}", "\n\n", t)
    t = re.sub(r"[ \t]{2,}", " ", t)
    return t.strip()


def _split_paragraphs(text: str) -> List[str]:
    t = _normalize_text(text)
    if not t:
        return []
    return [p.strip() for p in re.split(r"\n\s*\n", t, flags=re.MULTILINE) if p.strip()]


def _safe_filename(s: str, fallback: str = "olivetti") -> str:
    s = re.sub(r"[^\w\- ]+", "", (s or "").strip()).strip()
    s = re.sub(r"\s+", "_", s)
    return s[:80] if s else fallback


def _clamp_text(s: str, max_chars: int = 12000) -> str:
    s = s or ""
    if len(s) <= max_chars:
        return s
    return s[: max_chars - 40] + "\n\n… (truncated) …"


# ============================================================
# VECTOR / VOICE VAULT (lightweight, no external deps)
# ============================================================
def _tokenize(text: str) -> List[str]:
    return [w.lower() for w in WORD_RE.findall(text or "")]


def _hash_vec(text: str, dims: int = 512) -> List[float]:
    vec = [0.0] * dims
    toks = _tokenize(text)
    for t in toks:
        h = int(hashlib.md5(t.encode("utf-8")).hexdigest(), 16)
        vec[h % dims] += 1.0
    for i, v in enumerate(vec):
        if v > 0:
            vec[i] = 1.0 + math.log(v)
    return vec


def _cosine(a: List[float], b: List[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(y * y for y in b))
    if na == 0 or nb == 0:
        return 0.0
    return dot / (na * nb)


def default_voice_vault() -> Dict[str, Any]:
    ts = now_ts()
    return {
        "Voice A": {"created_ts": ts, "lanes": {ln: [] for ln in LANES}},
        "Voice B": {"created_ts": ts, "lanes": {ln: [] for ln in LANES}},
    }


def rebuild_vectors_in_voice_vault(compact_voices: Dict[str, Any]) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    for vname, v in (compact_voices or {}).items():
        created_ts = v.get("created_ts") or now_ts()
        lanes_in = v.get("lanes", {}) or {}
        lanes_out: Dict[str, Any] = {ln: [] for ln in LANES}
        for ln in LANES:
            samples = lanes_in.get(ln, []) or []
            for s in samples:
                txt = _normalize_text(s.get("text", ""))
                if not txt:
                    continue
                lanes_out[ln].append({
                    "ts": s.get("ts") or now_ts(),
                    "text": txt,
                    "vec": _hash_vec(txt),
                })
        out[vname] = {"created_ts": created_ts, "lanes": lanes_out}
    return out


def compact_voice_vault(voices: Dict[str, Any]) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    for vname, v in (voices or {}).items():
        lanes_out: Dict[str, Any] = {}
        for ln in LANES:
            samples = (v.get("lanes", {}) or {}).get(ln, []) or []
            lanes_out[ln] = [
                {"ts": s.get("ts"), "text": s.get("text", "")}
                for s in samples
                if (s.get("text") or "").strip()
            ]
        out[vname] = {"created_ts": v.get("created_ts"), "lanes": lanes_out}
    return out


def voice_names_for_selector() -> List[str]:
    base = ["— None —", "Voice A", "Voice B"]
    customs = sorted([k for k in (st.session_state.voices or {}).keys() if k not in ("Voice A", "Voice B")])
    return base + customs


def create_custom_voice(name: str) -> bool:
    n = (name or "").strip()
    if not n:
        return False
    if n in (st.session_state.voices or {}):
        return False
    st.session_state.voices[n] = {"created_ts": now_ts(), "lanes": {ln: [] for ln in LANES}}
    return True


def add_voice_sample(voice_name: str, lane: str, text: str) -> bool:
    vn = (voice_name or "").strip()
    if not vn:
        return False
    lane = lane if lane in LANES else "Narration"
    t = _normalize_text(text)
    if not t:
        return False
    v = (st.session_state.voices or {}).get(vn)
    if not v:
        # auto-create
        create_custom_voice(vn)
        v = st.session_state.voices.get(vn)
    v.setdefault("lanes", {ln: [] for ln in LANES})
    v["lanes"].setdefault(lane, [])
    v["lanes"][lane].append({"ts": now_ts(), "text": t, "vec": _hash_vec(t)})
    # cap per lane (keeps app fast)
    if len(v["lanes"][lane]) > 60:
        v["lanes"][lane] = v["lanes"][lane][-60:]
    st.session_state.voices[vn] = v
    return True


def delete_voice_sample(voice_name: str, lane: str, index_from_end: int = 0) -> bool:
    vn = (voice_name or "").strip()
    v = (st.session_state.voices or {}).get(vn)
    if not v:
        return False
    lane = lane if lane in LANES else "Narration"
    arr = (v.get("lanes", {}) or {}).get(lane, []) or []
    if not arr:
        return False
    idx = len(arr) - 1 - int(index_from_end)
    if idx < 0 or idx >= len(arr):
        return False
    arr.pop(idx)
    v["lanes"][lane] = arr
    st.session_state.voices[vn] = v
    return True


def analyze_style_samples(text: str) -> List[Dict[str, Any]]:
    """
    Analyze text for strongest writing samples based on:
    - Sentence variety and rhythm
    - Vocabulary richness (unique words)
    - Imagery and sensory language
    - Syntactic complexity
    Returns list of highlighted samples with scores.
    """
    if not text or not text.strip():
        return []
    
    # Split into sentences
    import re
    sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
    if not sentences:
        return []
    
    scored_samples = []
    
    for i, sent in enumerate(sentences):
        if len(sent) < 10:  # Skip very short sentences
            continue
            
        words = _tokenize(sent)
        if len(words) < 3:
            continue
        
        # Vocabulary richness (unique words ratio)
        unique_ratio = len(set(words)) / max(len(words), 1)
        
        # Sensory/imagery words
        sensory_words = {"saw", "see", "seen", "looked", "felt", "feel", "heard", "hear", "smelled", "smell", "tasted", "taste",
                        "touch", "sound", "sight", "color", "light", "dark", "bright", "shadow", "whisper", "shout",
                        "warm", "cold", "hot", "cool", "soft", "hard", "rough", "smooth", "bitter", "sweet", "sour"}
        sensory_count = sum(1 for w in words if w.lower() in sensory_words)
        
        # Sentence length variance (prefer medium-length sentences with complexity)
        length_score = 1.0 - abs(len(words) - 15) / 30.0  # Optimal around 15 words
        length_score = max(0.0, min(1.0, length_score))
        
        # Strong verbs (action)
        strong_verbs = ACTION_VERBS
        verb_count = sum(1 for w in words if w.lower() in strong_verbs)
        
        # Thought/interiority depth
        thought_count = sum(1 for w in words if w.lower() in THOUGHT_WORDS)
        
        # Calculate composite score
        score = (
            unique_ratio * 30.0 +          # Vocabulary richness
            sensory_count * 15.0 +         # Sensory language
            length_score * 20.0 +          # Optimal length
            verb_count * 10.0 +            # Action/dynamism
            thought_count * 5.0            # Depth
        )
        
        scored_samples.append({
            "text": sent,
            "score": score,
            "index": i,
            "words": len(words),
            "unique_ratio": unique_ratio,
            "sensory": sensory_count,
            "verbs": verb_count
        })
    
    # Sort by score and return top samples
    scored_samples.sort(key=lambda x: x["score"], reverse=True)
    return scored_samples[:10]  # Top 10 samples


def extract_entities(text: str) -> Dict[str, List[str]]:
    """
    Extract named entities from text for canon checking.
    Returns dict with entity types: characters, locations, dates, objects.
    """
    import re
    
    entities = {
        "characters": [],
        "locations": [],
        "dates": [],
        "objects": []
    }
    
    if not text or not text.strip():
        return entities
    
    # Find capitalized words (potential character names)
    capitalized = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', text)
    entities["characters"] = list(set(capitalized))[:20]  # Limit to top 20
    
    # Find day/month/time references
    date_patterns = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday',
                     'January', 'February', 'March', 'April', 'May', 'June', 'July', 'August',
                     'September', 'October', 'November', 'December']
    for pattern in date_patterns:
        if pattern in text:
            entities["dates"].append(pattern)
    
    return entities


def analyze_canon_conformity(draft_text: str) -> List[Dict[str, Any]]:
    """
    Analyze draft against Story Bible for continuity violations.
    Returns list of potential issues with confidence scores.
    """
    if not draft_text or not draft_text.strip():
        return []
    
    issues = []
    
    # Get Story Bible content
    synopsis = (st.session_state.synopsis or "").lower()
    characters_bible = (st.session_state.characters or "").lower()
    world_bible = (st.session_state.world or "").lower()
    outline_bible = (st.session_state.outline or "").lower()
    
    # If Story Bible is empty, no canon to check
    if not any([synopsis, characters_bible, world_bible, outline_bible]):
        return []
    
    # Split draft into paragraphs for analysis
    paragraphs = [p.strip() for p in draft_text.split('\n\n') if p.strip()]
    
    for para_idx, para in enumerate(paragraphs):
        para_lower = para.lower()
        words = _tokenize(para)
        
        # Check for character trait contradictions
        # Look for descriptions that might conflict
        if characters_bible:
            # Eye color check
            eye_colors = ['blue eyes', 'green eyes', 'brown eyes', 'gray eyes', 'grey eyes', 'hazel eyes', 'amber eyes']
            for color in eye_colors:
                if color in para_lower:
                    # Check if Story Bible says different
                    for other_color in eye_colors:
                        if other_color != color and other_color in characters_bible:
                            # Extract character name near the eye color
                            import re
                            match = re.search(r'([A-Z][a-z]+).*?' + color.replace(' ', r'\s+'), para, re.IGNORECASE)
                            if match:
                                char_name = match.group(1)
                                if char_name.lower() in characters_bible:
                                    issues.append({
                                        "type": "character_trait",
                                        "severity": "error",
                                        "confidence": 85,
                                        "paragraph_index": para_idx,
                                        "text_snippet": para[:100],
                                        "issue": f"'{char_name}' has {color} in draft, but Story Bible suggests {other_color}",
                                        "resolution_options": ["Update Story Bible", "Fix Draft", "Ignore"]
                                    })
        
        # Check for dead character mentions
        death_markers = ['died', 'dead', 'killed', 'perished', 'deceased']
        for marker in death_markers:
            if marker in para_lower:
                # Find character name near death marker
                import re
                match = re.search(r'([A-Z][a-z]+).*?' + marker, para, re.IGNORECASE)
                if match:
                    dead_char = match.group(1).lower()
                    # Check if this character appears later in draft
                    later_paras = paragraphs[para_idx + 1:]
                    for later_idx, later_para in enumerate(later_paras, start=para_idx + 1):
                        if dead_char in later_para.lower():
                            # Check if they're doing living things
                            living_verbs = ['said', 'walked', 'ran', 'thought', 'smiled', 'laughed', 'grabbed']
                            if any(verb in later_para.lower() for verb in living_verbs):
                                issues.append({
                                    "type": "continuity",
                                    "severity": "error",
                                    "confidence": 75,
                                    "paragraph_index": later_idx,
                                    "text_snippet": later_para[:100],
                                    "issue": f"'{match.group(1)}' appears active after being marked as {marker}",
                                    "resolution_options": ["Fix Draft", "Ignore"]
                                })
                                break
        
        # Check for timeline contradictions
        if outline_bible:
            days = ['monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday']
            for day in days:
                if day in para_lower:
                    # Check if outline specifies different day for this event
                    # Look for event keywords in paragraph
                    event_keywords = ['heist', 'meeting', 'battle', 'ceremony', 'wedding', 'funeral', 'attack']
                    for event in event_keywords:
                        if event in para_lower:
                            # Check if outline mentions this event with different day
                            for other_day in days:
                                if other_day != day and event in outline_bible and other_day in outline_bible:
                                    issues.append({
                                        "type": "timeline",
                                        "severity": "warning",
                                        "confidence": 60,
                                        "paragraph_index": para_idx,
                                        "text_snippet": para[:100],
                                        "issue": f"Event '{event}' on {day.title()}, but outline may indicate {other_day.title()}",
                                        "resolution_options": ["Update Outline", "Fix Draft", "Ignore"]
                                    })
        
        # Check for world-building contradictions
        if world_bible:
            # Technology level check
            modern_tech = ['phone', 'computer', 'car', 'internet', 'email', 'television', 'airplane']
            medieval_tech = ['sword', 'horse', 'castle', 'knight', 'dragon', 'magic', 'spell']
            
            has_modern = any(tech in para_lower for tech in modern_tech)
            has_medieval = any(tech in para_lower for tech in medieval_tech)
            
            if has_modern and has_medieval:
                if 'medieval' in world_bible or 'fantasy' in world_bible:
                    issues.append({
                        "type": "world_building",
                        "severity": "warning",
                        "confidence": 70,
                        "paragraph_index": para_idx,
                        "text_snippet": para[:100],
                        "issue": "Modern technology in medieval/fantasy setting (Story Bible suggests period setting)",
                        "resolution_options": ["Update World", "Fix Draft", "Ignore"]
                    })
    
    # Filter out ignored issues
    ignored_flags = st.session_state.get("canon_ignored_flags", [])
    filtered_issues = []
    for issue in issues:
        flag_id = f"{issue['type']}_{issue['paragraph_index']}_{issue['issue'][:30]}"
        if flag_id not in ignored_flags:
            filtered_issues.append(issue)
    
    return filtered_issues


def analyze_voice_conformity(text: str) -> List[Dict[str, Any]]:
    """
    Analyze how well text conforms to active Voice Bible controls.
    Returns list of paragraphs with conformity scores (0-100).
    Higher score = better conformity, lower = more deviation.
    """
    if not text or not text.strip():
        return []
    
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    if not paragraphs:
        return []
    
    results = []
    
    for idx, para in enumerate(paragraphs):
        words = _tokenize(para)
        if len(words) < 3:
            continue
        
        score = 100.0  # Start at perfect conformity
        issues = []
        
        # Check POV conformity (if technical controls enabled)
        if st.session_state.vb_technical_on:
            pov = st.session_state.pov
            first_person = sum(1 for w in words if w.lower() in ['i', 'me', 'my', 'mine', 'myself', 'we', 'us', 'our'])
            third_person = sum(1 for w in words if w.lower() in ['he', 'she', 'they', 'him', 'her', 'them', 'his', 'hers', 'their'])
            
            if pov == "First" and third_person > first_person:
                score -= 20
                issues.append("POV mismatch: too much third-person")
            elif pov in ["Close Third", "Omniscient"] and first_person > third_person * 0.5:
                score -= 20
                issues.append("POV mismatch: too much first-person")
        
        # Check tense conformity (if technical controls enabled)
        if st.session_state.vb_technical_on:
            tense = st.session_state.tense
            past_markers = sum(1 for w in words if w.endswith('ed') or w in ['was', 'were', 'had', 'did'])
            present_markers = sum(1 for w in words if w.endswith('s') and len(w) > 2 or w in ['is', 'are', 'am', 'does', 'do'])
            
            if tense == "Past" and present_markers > past_markers:
                score -= 15
                issues.append("Tense mismatch: too much present tense")
            elif tense == "Present" and past_markers > present_markers:
                score -= 15
                issues.append("Tense mismatch: too much past tense")
        
        # Check style conformity (if style engine enabled)
        if st.session_state.vb_style_on:
            style = st.session_state.writing_style
            
            # LYRICAL expects poetic/sensory language
            if style == "LYRICAL":
                sensory_count = sum(1 for w in words if w.lower() in {
                    "light", "dark", "shadow", "sound", "whisper", "touch", "soft", "rough",
                    "color", "bright", "warm", "cold", "scent", "taste"
                })
                if sensory_count < len(words) * 0.03:  # Less than 3% sensory
                    score -= 15
                    issues.append("Style: needs more sensory/poetic language")
            
            # SPARSE expects brevity
            elif style == "SPARSE":
                avg_word_len = sum(len(w) for w in words) / max(len(words), 1)
                if avg_word_len > 5.5:
                    score -= 15
                    issues.append("Style: words too long for sparse style")
            
            # ORNATE expects complexity
            elif style == "ORNATE":
                avg_word_len = sum(len(w) for w in words) / max(len(words), 1)
                if avg_word_len < 4.5:
                    score -= 15
                    issues.append("Style: needs more elaborate vocabulary")
        
        # Check genre conformity (if genre intelligence enabled)
        if st.session_state.vb_genre_on:
            genre = st.session_state.genre
            
            # NOIR expects hardboiled tone
            if genre == "Noir":
                noir_words = sum(1 for w in words if w.lower() in {
                    "dark", "shadow", "night", "smoke", "rain", "gun", "blood", "dead", "dame"
                })
                if noir_words < 1 and len(words) > 30:
                    score -= 10
                    issues.append("Genre: lacks noir atmosphere")
            
            # HORROR expects tension
            elif genre == "Horror":
                horror_words = sum(1 for w in words if w.lower() in {
                    "fear", "terror", "scream", "blood", "dark", "shadow", "death", "cold", "alone"
                })
                if horror_words < 1 and len(words) > 30:
                    score -= 10
                    issues.append("Genre: lacks horror elements")
        
        # Check voice lock violations (if enabled)
        if st.session_state.vb_lock_on and st.session_state.voice_lock_prompt:
            lock_prompt = st.session_state.voice_lock_prompt.lower()
            if "no adverb" in lock_prompt or "never adverb" in lock_prompt:
                adverbs = sum(1 for w in words if w.endswith('ly'))
                if adverbs > 0:
                    score -= 25
                    issues.append(f"Voice Lock: {adverbs} adverb(s) found")
            if "no passive" in lock_prompt or "never passive" in lock_prompt:
                passive = sum(1 for w in words if w in ['was', 'were', 'been', 'being'])
                if passive > 0:
                    score -= 20
                    issues.append("Voice Lock: passive voice detected")
        
        # Ensure score stays in bounds
        score = max(0.0, min(100.0, score))
        
        results.append({
            "text": para,
            "score": score,
            "index": idx,
            "words": len(words),
            "issues": issues
        })
    
    return results


def retrieve_exemplars(voice_name: str, lane: str, query_text: str, k: int = 3) -> List[str]:
    v = (st.session_state.voices or {}).get(voice_name)
    if not v:
        return []
    lane = lane if lane in LANES else "Narration"
    pool = v.get("lanes", {}).get(lane, []) or []
    if not pool:
        return []
    qv = _hash_vec(query_text)
    scored = [(_cosine(qv, s.get("vec", [])), s.get("text", "")) for s in pool[-140:]]
    scored.sort(key=lambda x: x[0], reverse=True)
    return [txt for score, txt in scored[:k] if score > 0.0 and txt][:k]


def retrieve_mixed_exemplars(voice_name: str, lane: str, query_text: str) -> List[str]:
    lane_ex = retrieve_exemplars(voice_name, lane, query_text, k=2)
    if lane == "Narration":
        return lane_ex if lane_ex else retrieve_exemplars(voice_name, "Narration", query_text, k=3)
    nar_ex = retrieve_exemplars(voice_name, "Narration", query_text, k=1)
    out = lane_ex + [x for x in nar_ex if x not in lane_ex]
    return out[:3]


# ============================================================
# LANE DETECTION (lightweight)
# ============================================================
THOUGHT_WORDS = {
    "think",
    "thought",
    "felt",
    "wondered",
    "realized",
    "remembered",
    "knew",
    "noticed",
    "decided",
    "hoped",
    "feared",
    "wanted",
    "imagined",
    "could",
    "should",
    "would",
}
ACTION_VERBS = {
    "run",
    "ran",
    "walk",
    "walked",
    "grab",
    "grabbed",
    "push",
    "pushed",
    "pull",
    "pulled",
    "slam",
    "slammed",
    "hit",
    "struck",
    "kick",
    "kicked",
    "turn",
    "turned",
    "snap",
    "snapped",
    "dive",
    "dived",
    "duck",
    "ducked",
    "rush",
    "rushed",
    "lunge",
    "lunged",
    "climb",
    "climbed",
    "drop",
    "dropped",
    "throw",
    "threw",
    "fire",
    "fired",
    "aim",
    "aimed",
    "break",
    "broke",
    "shatter",
    "shattered",
    "step",
    "stepped",
    "move",
    "moved",
    "reach",
    "reached",
}


def detect_lane(paragraph: str) -> str:
    p = (paragraph or "").strip()
    if not p:
        return "Narration"

    quote_count = p.count('"') + p.count("“") + p.count("”")
    has_dialogue_punct = p.startswith(("—", "- ", "“", '"'))

    dialogue_score = 0.0
    if quote_count >= 2:
        dialogue_score += 2.5
    if has_dialogue_punct:
        dialogue_score += 1.5

    toks = _tokenize(p)
    interior_score = 0.0
    if toks:
        first_person = sum(1 for t in toks if t in ("i", "me", "my", "mine", "myself"))
        thought_hits = sum(1 for t in toks if t in THOUGHT_WORDS)
        if first_person >= 2 and thought_hits >= 1:
            interior_score += 2.2

    action_score = 0.0
    if toks:
        verb_hits = sum(1 for t in toks if t in ACTION_VERBS)
        if verb_hits >= 2:
            action_score += 1.6
        if "!" in p:
            action_score += 0.3

    scores = {"Dialogue": dialogue_score, "Interiority": interior_score, "Action": action_score, "Narration": 0.25}
    lane = max(scores.items(), key=lambda kv: kv[1])[0]
    return "Narration" if scores[lane] < 0.9 else lane


def current_lane_from_draft(text: str) -> str:
    paras = _split_paragraphs(text)
    if not paras:
        return "Narration"
    return detect_lane(paras[-1])


# ============================================================
# INTENSITY (GLOBAL AI AGGRESSION KNOB)
# ============================================================
def intensity_profile(x: float) -> str:
    if x <= 0.25:
        return "LOW: conservative, literal, minimal invention, prioritize continuity and clarity."
    if x <= 0.60:
        return "MED: balanced creativity, stronger phrasing, controlled invention within canon."
    if x <= 0.85:
        return "HIGH: bolder choices, richer specificity; still obey canon and lane."
    return "MAX: aggressive originality and voice; still obey canon, no derailments."


def temperature_from_intensity(x: float) -> float:
    x = max(0.0, min(1.0, float(x)))
    return 0.15 + (x * 0.95)


def _get_voice_bible_summary() -> str:
    """Generate a brief summary of active Voice Bible controls for status display."""
    parts = []
    ai_int = float(st.session_state.ai_intensity)
    
    # AI Intensity is always shown
    if ai_int <= 0.25:
        parts.append("AI:LOW")
    elif ai_int <= 0.60:
        parts.append("AI:MED")
    elif ai_int <= 0.85:
        parts.append("AI:HIGH")
    else:
        parts.append("AI:MAX")
    
    # Show active Voice Bible controls
    if st.session_state.vb_style_on:
        parts.append(f"Style:{st.session_state.writing_style}")
    if st.session_state.vb_genre_on:
        parts.append(f"Genre:{st.session_state.genre}")
    if st.session_state.vb_trained_on and st.session_state.trained_voice != "— None —":
        parts.append(f"Voice:{st.session_state.trained_voice}")
    if st.session_state.vb_match_on:
        parts.append("Match:ON")
    if st.session_state.vb_lock_on:
        parts.append("🔒Lock:ON")
    if st.session_state.vb_technical_on:
        parts.append(f"Tech:{st.session_state.pov}/{st.session_state.tense}")
    
    return " • ".join(parts) if len(parts) > 1 else parts[0] if parts else "Defaults"


def _verify_system_integrity() -> Dict[str, bool]:
    """
    ═══════════════════════════════════════════════════════════════
    SYSTEM INTEGRITY CHECK - Verifies all components are connected.
    Run at startup to ensure the spiderweb is fully wired.
    ═══════════════════════════════════════════════════════════════
    """
    checks = {
        "ai_intensity_exists": "ai_intensity" in st.session_state,
        "story_bible_sections": all(k in st.session_state for k in ["synopsis", "genre_style_notes", "world", "characters", "outline"]),
        "voice_bible_controls": all(k in st.session_state for k in ["vb_style_on", "vb_genre_on", "vb_trained_on", "vb_match_on", "vb_lock_on", "vb_technical_on"]),
        "style_banks_exist": "style_banks" in st.session_state,
        "voices_exist": "voices" in st.session_state,
        "project_system": all(k in st.session_state for k in ["projects", "active_bay", "project_id"]),
        "functions_callable": all(callable(f) for f in [
            build_partner_brief, 
            call_openai, 
            partner_action, 
            retrieve_style_exemplars,
            retrieve_mixed_exemplars,
            generate_story_bible_section,
            temperature_from_intensity
        ]),
    }
    return checks


# ============================================================
# PROJECT MODEL
# ============================================================
def _fingerprint_story_bible(sb: Dict[str, str]) -> str:
    parts = [
        (sb.get("synopsis", "") or "").strip(),
        (sb.get("genre_style_notes", "") or "").strip(),
        (sb.get("world", "") or "").strip(),
        (sb.get("characters", "") or "").strip(),
        (sb.get("outline", "") or "").strip(),
    ]
    blob = "\n\n---\n\n".join(parts)
    return hashlib.md5(blob.encode("utf-8")).hexdigest()


def new_project_payload(title: str) -> Dict[str, Any]:
    ts = now_ts()
    title = title.strip() if title.strip() else "Untitled Project"
    story_bible_id = hashlib.md5(f"sb|{title}|{ts}".encode("utf-8")).hexdigest()[:12]
    return {
        "id": hashlib.md5(f"{title}|{ts}".encode("utf-8")).hexdigest()[:12],
        "title": title,
        "created_ts": ts,
        "updated_ts": ts,
        "bay": "NEW",
        "draft": "",
        "story_bible_id": story_bible_id,
        "story_bible_created_ts": ts,
        "story_bible_fingerprint": "",
        "story_bible": {"synopsis": "", "genre_style_notes": "", "world": "", "characters": "", "outline": ""},
        "voice_bible": {
            "vb_style_on": True,
            "vb_genre_on": True,
            "vb_trained_on": False,
            "vb_match_on": False,
            "vb_lock_on": False,
            "vb_technical_on": True,
            "writing_style": "Neutral",
            "genre": "Literary",
            "trained_voice": "— None —",
            "voice_sample": "",
            "voice_lock_prompt": "",
            "style_intensity": 0.6,
            "genre_intensity": 0.6,
            "trained_intensity": 0.7,
            "match_intensity": 0.8,
            "lock_intensity": 1.0,
            "technical_intensity": 0.8,
            "pov": "Close Third",
            "tense": "Past",
            "ai_intensity": 0.75,
        },
        # locks removed: Story Bible is always editable
        "voices": default_voice_vault(),
        "style_banks": default_style_banks(),
    }


# ============================================================

# ============================================================
# ENGINE STYLE BANKS (project/workspace) — trainable writing styles
# ============================================================
def default_style_banks() -> Dict[str, Any]:
    ts = now_ts()
    return {s: {"created_ts": ts, "lanes": {ln: [] for ln in LANES}} for s in ENGINE_STYLES}


def rebuild_vectors_in_style_banks(banks: Dict[str, Any]) -> Dict[str, Any]:
    src = banks or {}
    out: Dict[str, Any] = {}
    for style in ENGINE_STYLES:
        b = (src.get(style) or {}) if isinstance(src, dict) else {}
        lanes = b.get("lanes") or {}
        new_lanes: Dict[str, List[Dict[str, Any]]] = {}
        for ln in LANES:
            samples = (lanes.get(ln) or []) if isinstance(lanes, dict) else []
            rebuilt: List[Dict[str, Any]] = []
            for it in samples:
                if isinstance(it, str):
                    t = it.strip()
                    if not t:
                        continue
                    rebuilt.append({"ts": now_ts(), "text": t, "vec": _hash_vec(t)})
                    continue
                if not isinstance(it, dict):
                    continue
                t = (it.get("text") or "").strip()
                if not t:
                    continue
                vec = it.get("vec") if isinstance(it.get("vec"), list) else None
                if not vec:
                    vec = _hash_vec(t)
                rebuilt.append({"ts": it.get("ts") or now_ts(), "text": t, "vec": vec})
            new_lanes[ln] = rebuilt
        out[style] = {"created_ts": b.get("created_ts") or now_ts(), "lanes": new_lanes}
    return out if out else default_style_banks()


def compact_style_banks(banks: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(banks, dict):
        return default_style_banks()
    out: Dict[str, Any] = {}
    for style in ENGINE_STYLES:
        b = banks.get(style) or {}
        lanes = b.get("lanes") or {}
        c_lanes: Dict[str, List[Dict[str, Any]]] = {}
        for ln in LANES:
            ss = (lanes.get(ln) or []) if isinstance(lanes, dict) else []
            cleaned: List[Dict[str, Any]] = []
            for it in ss:
                if not isinstance(it, dict):
                    continue
                t = (it.get("text") or "").strip()
                if not t:
                    continue
                cleaned.append({"ts": it.get("ts") or now_ts(), "text": _clamp_text(t, 9000)})
            c_lanes[ln] = cleaned
        out[style] = {"created_ts": b.get("created_ts") or now_ts(), "lanes": c_lanes}
    return out


def add_style_samples(style: str, lane: str, raw_text: str, split_mode: str = "Paragraphs", cap_per_lane: int = 250) -> int:
    style = (style or "").strip().upper()
    if style not in ENGINE_STYLES:
        return 0
    lane = lane if lane in LANES else "Narration"
    t = _normalize_text(raw_text)
    if not t.strip():
        return 0

    parts = _split_paragraphs(t) if split_mode == "Paragraphs" else [t.strip()]
    parts = [p for p in parts if len(p.strip()) >= 40]

    sb = st.session_state.get("style_banks")
    if not isinstance(sb, dict) or style not in sb:
        sb = rebuild_vectors_in_style_banks(default_style_banks())
        st.session_state.style_banks = sb

    bank = sb.get(style) or {}
    lanes = bank.get("lanes") or {}
    lane_list = list((lanes.get(lane) or [])) if isinstance(lanes, dict) else []

    added = 0
    for p in parts:
        p = _clamp_text(p.strip(), 9000)
        lane_list.append({"ts": now_ts(), "text": p, "vec": _hash_vec(p)})
        added += 1

    # cap: keep newest
    lane_list = lane_list[-cap_per_lane:]
    lanes[lane] = lane_list
    bank["lanes"] = lanes
    sb[style] = bank
    st.session_state.style_banks = sb
    return added


def delete_last_style_sample(style: str, lane: str) -> bool:
    style = (style or "").strip().upper()
    if style not in ENGINE_STYLES:
        return False
    lane = lane if lane in LANES else "Narration"
    sb = st.session_state.get("style_banks") or {}
    bank = (sb.get(style) or {})
    lanes = bank.get("lanes") or {}
    lane_list = (lanes.get(lane) or [])
    if not lane_list:
        return False
    lane_list.pop()
    lanes[lane] = lane_list
    bank["lanes"] = lanes
    sb[style] = bank
    st.session_state.style_banks = sb
    return True


def clear_style_lane(style: str, lane: str) -> None:
    style = (style or "").strip().upper()
    if style not in ENGINE_STYLES:
        return
    lane = lane if lane in LANES else "Narration"
    sb = st.session_state.get("style_banks") or rebuild_vectors_in_style_banks(default_style_banks())
    bank = sb.get(style) or {}
    lanes = bank.get("lanes") or {}
    lanes[lane] = []
    bank["lanes"] = lanes
    sb[style] = bank
    st.session_state.style_banks = sb


def retrieve_style_exemplars(style: str, lane: str, query: str, k: int = 2) -> List[str]:
    style = (style or "").strip().upper()
    if style not in ENGINE_STYLES:
        return []
    lane = lane if lane in LANES else "Narration"
    sb = st.session_state.get("style_banks") or {}
    bank = sb.get(style) or {}
    lanes = bank.get("lanes") or {}
    pool = (lanes.get(lane) or [])
    if not pool:
        # fallback: all lanes pooled
        pool = []
        for ln in LANES:
            pool.extend(lanes.get(ln) or [])
    if not pool:
        return []
    # favor newest slice for speed
    pool = pool[-160:]
    qv = _hash_vec(query or "")
    scored = []
    for it in pool:
        if not isinstance(it, dict):
            continue
        vec = it.get("vec")
        if not isinstance(vec, list):
            continue
        scored.append((_cosine(qv, vec), it.get("text") or ""))
    scored.sort(key=lambda x: x[0], reverse=True)
    out = [t.strip() for _, t in scored[: max(0, k)] if (t or "").strip()]
    return out


_ENGINE_STYLE_GUIDE = {
    "NARRATIVE": "Narrative clarity, clean cause→effect, confident pacing. Prioritize story logic and readability.",
    "DESCRIPTIVE": "Sensory precision, spatial clarity, vivid concrete nouns, controlled detail density (no purple bloat).",
    "EMOTIONAL": "Interior depth, subtext, emotional specificity. Show the feeling through behavior, sensation, and thought.",
    "LYRICAL": "Rhythm, musical syntax, image-forward language, elegant metaphor with restraint. Make prose sing without obscuring meaning.",
}


def engine_style_directive(style: str, intensity: float, lane: str) -> str:
    style = (style or "").strip().upper()
    base = _ENGINE_STYLE_GUIDE.get(style, "")
    x = float(intensity)
    if x <= 0.33:
        mod = "Keep it subtle and controlled. Minimal overt stylization."
    elif x <= 0.66:
        mod = "Medium stylization. Let the style clearly shape diction and cadence."
    else:
        mod = "High stylization. Strong stylistic fingerprint, but still professional and coherent."
    return f"{base}\\nLane: {lane}\\n{mod}"


# STORY BIBLE WORKSPACE (pre-project creation space)
# ============================================================
def default_story_bible_workspace() -> Dict[str, Any]:
    ts = now_ts()
    return {
        "workspace_story_bible_id": hashlib.md5(f"wsb|{ts}".encode("utf-8")).hexdigest()[:12],
        "workspace_story_bible_created_ts": ts,
        "title": "",
        "draft": "",
        "story_bible": {"synopsis": "", "genre_style_notes": "", "world": "", "characters": "", "outline": ""},
        "voice_sample": "",
        "ai_intensity": 0.75,
        "voices": default_voice_vault(),
        "style_banks": default_style_banks(),
    }


def in_workspace_mode() -> bool:
    return (st.session_state.active_bay == "NEW") and (st.session_state.project_id is None)


def save_workspace_from_session() -> None:
    w = st.session_state.sb_workspace or default_story_bible_workspace()
    w["title"] = st.session_state.get("workspace_title", w.get("title", ""))
    w["draft"] = st.session_state.main_text
    w["story_bible"] = {
        "synopsis": st.session_state.synopsis,
        "genre_style_notes": st.session_state.genre_style_notes,
        "world": st.session_state.world,
        "characters": st.session_state.characters,
        "outline": st.session_state.outline,
    }
    w["voice_sample"] = st.session_state.voice_sample
    w["ai_intensity"] = float(st.session_state.ai_intensity)
    w["voices"] = compact_voice_vault(st.session_state.voices)
    w["style_banks"] = compact_style_banks(st.session_state.get("style_banks") or rebuild_vectors_in_style_banks(default_style_banks()))
    st.session_state.sb_workspace = w


def set_ai_intensity(val: float) -> None:
    """Safely set ai_intensity without mutating widget-bound state after creation."""
    val = float(val)
    if "ai_intensity" not in st.session_state:
        st.session_state["ai_intensity"] = val
    else:
        st.session_state["ai_intensity_pending"] = val
        st.session_state["_apply_pending_widget_state"] = True


def load_workspace_into_session() -> None:
    w = st.session_state.sb_workspace or default_story_bible_workspace()
    sb = w.get("story_bible", {}) or {}
    st.session_state.project_id = None
    st.session_state.project_title = "—"
    st.session_state.main_text = w.get("draft", "") or ""
    st.session_state.synopsis = sb.get("synopsis", "") or ""
    st.session_state.genre_style_notes = sb.get("genre_style_notes", "") or ""
    st.session_state.world = sb.get("world", "") or ""
    st.session_state.characters = sb.get("characters", "") or ""
    st.session_state.outline = sb.get("outline", "") or ""
    st.session_state.voice_sample = w.get("voice_sample", "") or ""
    set_ai_intensity(float(w.get("ai_intensity", 0.75)))
    st.session_state.voices = rebuild_vectors_in_voice_vault(w.get("voices", default_voice_vault()))
    st.session_state.voices_seeded = True
    st.session_state.style_banks = rebuild_vectors_in_style_banks(w.get("style_banks", default_style_banks()))
    st.session_state.workspace_title = w.get("title", "") or ""


def reset_workspace_story_bible(keep_templates: bool = True) -> None:
    old = st.session_state.sb_workspace or default_story_bible_workspace()
    neww = default_story_bible_workspace()
    if keep_templates:
        neww["voice_sample"] = old.get("voice_sample", "")
        neww["ai_intensity"] = float(old.get("ai_intensity", 0.75))
        neww["voices"] = old.get("voices", default_voice_vault())
        neww["style_banks"] = old.get("style_banks", default_style_banks())
    st.session_state.sb_workspace = neww
    if in_workspace_mode():
        load_workspace_into_session()


# ============================================================
# SESSION INIT
# ============================================================
def init_state() -> None:
    defaults: Dict[str, Any] = {
        "active_bay": "NEW",
        "projects": {},
        "active_project_by_bay": {b: None for b in BAYS},
        "sb_workspace": default_story_bible_workspace(),
        "workspace_title": "",
        "project_id": None,
        "project_title": "—",
        "autosave_time": None,
        "last_action": "—",
        "voice_status": "—",
        "main_text": "",
        "synopsis": "",
        "genre_style_notes": "",
        "world": "",
        "characters": "",
        "outline": "",
        "junk": "",
        "tool_output": "",
        "pending_action": None,
        "vb_style_on": True,
        "vb_genre_on": True,
        "vb_trained_on": False,
        "vb_match_on": False,
        "vb_lock_on": False,
        "vb_technical_on": True,
        "writing_style": "Neutral",
        "genre": "Literary",
        "trained_voice": "— None —",
        "voice_sample": "",
        "voice_lock_prompt": "",
        "style_intensity": 0.6,
        "genre_intensity": 0.6,
        "trained_intensity": 0.7,
        "match_intensity": 0.8,
        "lock_intensity": 1.0,
        "technical_intensity": 0.8,
        "pov": "Close Third",
        "tense": "Past",
        "ai_intensity": 0.75,
        "locks": {
            "story_bible_lock": True,
            "sb_edit_unlocked": False,
            "voice_fingerprint_lock": True,
            "lane_lock": False,
            "forced_lane": "Narration",
        },
        "voices": {},
        "voices_seeded": False,
        "style_banks": rebuild_vectors_in_style_banks(default_style_banks()),
        "last_saved_digest": "",
        "analyzed_style_samples": [],
        "voice_heatmap_data": [],
        "show_voice_heatmap": False,
        "canon_guardian_on": False,
        "canon_issues": [],
        "canon_ignored_flags": [],

        # internal UI helpers (not widgets)
        "ui_notice": "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


init_state()


# Apply pending widget values BEFORE widgets are created
if st.session_state.pop("_apply_pending_widget_state", False):
    if "ai_intensity_pending" in st.session_state:
        st.session_state["ai_intensity"] = float(st.session_state.pop("ai_intensity_pending"))



# ============================================================
# PROJECT <-> SESSION SYNC
# ============================================================
def load_project_into_session(pid: str) -> None:
    p = (st.session_state.projects or {}).get(pid)
    if not p:
        return

    st.session_state.project_id = pid
    st.session_state.project_title = p.get("title", "Untitled Project")
    st.session_state.main_text = p.get("draft", "")

    sb = p.get("story_bible", {}) or {}
    st.session_state.synopsis = sb.get("synopsis", "")
    st.session_state.genre_style_notes = sb.get("genre_style_notes", "")
    st.session_state.world = sb.get("world", "")
    st.session_state.characters = sb.get("characters", "")
    st.session_state.outline = sb.get("outline", "")

    vb = p.get("voice_bible", {}) or {}
    for k in [
        "vb_style_on",
        "vb_genre_on",
        "vb_trained_on",
        "vb_match_on",
        "vb_lock_on",
        "vb_technical_on",
        "writing_style",
        "genre",
        "trained_voice",
        "voice_sample",
        "voice_lock_prompt",
        "style_intensity",
        "genre_intensity",
        "trained_intensity",
        "match_intensity",
        "lock_intensity",
        "technical_intensity",
        "pov",
        "tense",
        "ai_intensity",
    ]:
        if k in vb:
            st.session_state[k] = vb[k]

    # locks removed: Story Bible is always editable

    st.session_state.voices = rebuild_vectors_in_voice_vault(p.get("voices", default_voice_vault()))
    st.session_state.voices_seeded = True


def save_session_into_project() -> None:
    pid = st.session_state.project_id
    if not pid:
        return
    p = (st.session_state.projects or {}).get(pid)
    if not p:
        return

    p["updated_ts"] = now_ts()
    p["draft"] = st.session_state.main_text
    p["story_bible"] = {
        "synopsis": st.session_state.synopsis,
        "genre_style_notes": st.session_state.genre_style_notes,
        "world": st.session_state.world,
        "characters": st.session_state.characters,
        "outline": st.session_state.outline,
    }
    p["voice_bible"] = {
        "vb_style_on": st.session_state.vb_style_on,
        "vb_genre_on": st.session_state.vb_genre_on,
        "vb_trained_on": st.session_state.vb_trained_on,
        "vb_match_on": st.session_state.vb_match_on,
        "vb_lock_on": st.session_state.vb_lock_on,
        "vb_technical_on": st.session_state.vb_technical_on,
        "writing_style": st.session_state.writing_style,
        "genre": st.session_state.genre,
        "trained_voice": st.session_state.trained_voice,
        "voice_sample": st.session_state.voice_sample,
        "voice_lock_prompt": st.session_state.voice_lock_prompt,
        "style_intensity": st.session_state.style_intensity,
        "genre_intensity": st.session_state.genre_intensity,
        "trained_intensity": st.session_state.trained_intensity,
        "match_intensity": st.session_state.match_intensity,
        "lock_intensity": st.session_state.lock_intensity,
        "technical_intensity": st.session_state.technical_intensity,
        "pov": st.session_state.pov,
        "tense": st.session_state.tense,
        "ai_intensity": float(st.session_state.ai_intensity),
    }
    # locks removed: Story Bible is always editable
    p["voices"] = compact_voice_vault(st.session_state.voices)
    p["style_banks"] = compact_style_banks(st.session_state.get("style_banks") or rebuild_vectors_in_style_banks(default_style_banks()))
    # keep fingerprint up to date
    try:
        p["story_bible_fingerprint"] = _fingerprint_story_bible(p["story_bible"])
    except Exception:
        pass


def list_projects_in_bay(bay: str) -> List[Tuple[str, str]]:
    items: List[Tuple[str, str]] = []
    for pid, p in (st.session_state.projects or {}).items():
        if isinstance(p, dict) and p.get("bay") == bay:
            items.append((pid, p.get("title", "Untitled")))
    items.sort(key=lambda x: (x[1] or "").lower())
    return items


def ensure_bay_has_active_project(bay: str) -> None:
    pid = (st.session_state.active_project_by_bay or {}).get(bay)
    if pid and pid in (st.session_state.projects or {}) and (st.session_state.projects[pid].get("bay") == bay):
        return
    items = list_projects_in_bay(bay)
    st.session_state.active_project_by_bay[bay] = items[0][0] if items else None


def switch_bay(target_bay: str) -> None:
    if in_workspace_mode():
        save_workspace_from_session()
    else:
        save_session_into_project()

    st.session_state.active_bay = target_bay
    ensure_bay_has_active_project(target_bay)
    pid = st.session_state.active_project_by_bay.get(target_bay)

    if pid:
        load_project_into_session(pid)
        st.session_state.voice_status = f"{target_bay}: {st.session_state.project_title}"
    else:
        st.session_state.project_id = None
        st.session_state.project_title = "—"
        if target_bay == "NEW":
            load_workspace_into_session()
            st.session_state.voice_status = "NEW: (Story Bible workspace)"
        else:
            st.session_state.main_text = ""
            st.session_state.synopsis = ""
            st.session_state.genre_style_notes = ""
            st.session_state.world = ""
            st.session_state.characters = ""
            st.session_state.outline = ""
            st.session_state.voice_sample = ""
            set_ai_intensity(0.75)
            st.session_state.voices = rebuild_vectors_in_voice_vault(default_voice_vault())
            st.session_state.voices_seeded = True
            st.session_state.voice_status = f"{target_bay}: (empty)"

    st.session_state.last_action = f"Bay → {target_bay}"


def create_project_from_current_bible(title: str) -> str:
    title = title.strip() if title.strip() else f"New Project {now_ts()}"
    source = "workspace" if in_workspace_mode() else "clone"

    source_story_bible_id = None
    source_story_bible_created_ts = None
    if source == "workspace":
        w = st.session_state.sb_workspace or default_story_bible_workspace()
        source_story_bible_id = w.get("workspace_story_bible_id")
        source_story_bible_created_ts = w.get("workspace_story_bible_created_ts")
    else:
        pid = st.session_state.project_id
        if pid and pid in (st.session_state.projects or {}):
            source_story_bible_id = st.session_state.projects[pid].get("story_bible_id")

    p = new_project_payload(title)
    p["bay"] = "NEW"
    p["draft"] = st.session_state.main_text
    p["story_bible"] = {
        "synopsis": st.session_state.synopsis,
        "genre_style_notes": st.session_state.genre_style_notes,
        "world": st.session_state.world,
        "characters": st.session_state.characters,
        "outline": st.session_state.outline,
    }

    if source == "workspace" and source_story_bible_id:
        p["story_bible_id"] = source_story_bible_id
        if source_story_bible_created_ts:
            p["story_bible_created_ts"] = source_story_bible_created_ts

    # story_bible_binding removed: Story Bible is always editable
    p["story_bible_fingerprint"] = _fingerprint_story_bible(p["story_bible"])

    # Voice templates + intensity
    p["voice_bible"]["voice_sample"] = st.session_state.voice_sample
    p["voice_bible"]["ai_intensity"] = float(st.session_state.ai_intensity)
    p["voices"] = compact_voice_vault(st.session_state.voices)
    p["style_banks"] = compact_style_banks(st.session_state.get("style_banks") or rebuild_vectors_in_style_banks(default_style_banks()))

    st.session_state.projects[p["id"]] = p
    st.session_state.active_project_by_bay["NEW"] = p["id"]

    if source == "workspace":
        reset_workspace_story_bible(keep_templates=True)

    return p["id"]


def promote_project(pid: str, to_bay: str) -> None:
    p = (st.session_state.projects or {}).get(pid)
    if not p:
        return
    p["bay"] = to_bay
    p["updated_ts"] = now_ts()


def next_bay(bay: str) -> Optional[str]:
    if bay == "NEW":
        return "ROUGH"
    if bay == "ROUGH":
        return "EDIT"
    if bay == "EDIT":
        return "FINAL"
    return None


# ============================================================
# AUTOSAVE (atomic + backup)
# ============================================================
def _payload() -> Dict[str, Any]:
    if in_workspace_mode():
        save_workspace_from_session()
    else:
        save_session_into_project()
    return {
        "meta": {"saved_at": now_ts(), "version": "olivetti-prod-stable-v1"},
        "active_bay": st.session_state.active_bay,
        "active_project_by_bay": st.session_state.active_project_by_bay,
        "sb_workspace": st.session_state.sb_workspace,
        "projects": st.session_state.projects,
    }


def _digest(payload: Dict[str, Any]) -> str:
    s = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    return hashlib.md5(s.encode("utf-8")).hexdigest()


def save_all_to_disk(force: bool = False) -> None:
    """Autosave state to disk with an atomic write and a simple backup."""
    try:
        os.makedirs(AUTOSAVE_DIR, exist_ok=True)
        payload = _payload()
        dig = _digest(payload)
        if (not force) and dig == st.session_state.last_saved_digest:
            return

        tmp_path = AUTOSAVE_PATH + ".tmp"
        bak_path = AUTOSAVE_PATH + ".bak"

        try:
            if os.path.exists(AUTOSAVE_PATH):
                import shutil

                shutil.copy2(AUTOSAVE_PATH, bak_path)
        except Exception:
            pass

        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        os.replace(tmp_path, AUTOSAVE_PATH)

        st.session_state.last_saved_digest = dig
    except Exception as e:
        st.session_state.voice_status = f"Autosave warning: {e}"


def load_all_from_disk() -> None:
    main_path = AUTOSAVE_PATH
    bak_path = AUTOSAVE_PATH + ".bak"

    def _boot_new() -> None:
        st.session_state.sb_workspace = st.session_state.get("sb_workspace") or default_story_bible_workspace()
        switch_bay("NEW")

    if (not os.path.exists(main_path)) and (not os.path.exists(bak_path)):
        _boot_new()
        return

    payload = None
    loaded_from = "primary"
    last_err = None
    for path, label in ((main_path, "primary"), (bak_path, "backup")):
        if not os.path.exists(path):
            continue
        try:
            with open(path, "r", encoding="utf-8") as f:
                payload = json.load(f)
            loaded_from = label
            break
        except Exception as e:
            last_err = e
            payload = None

    if payload is None:
        st.session_state.voice_status = f"Load warning: {last_err}"
        _boot_new()
        return

    try:
        projs = payload.get("projects", {})
        if isinstance(projs, dict):
            st.session_state.projects = projs

        apbb = payload.get("active_project_by_bay", {})
        if isinstance(apbb, dict):
            for b in BAYS:
                apbb.setdefault(b, None)
            st.session_state.active_project_by_bay = apbb

        w = payload.get("sb_workspace")
        if isinstance(w, dict) and w.get("workspace_story_bible_id"):
            st.session_state.sb_workspace = w
        else:
            st.session_state.sb_workspace = default_story_bible_workspace()

        # Migration guards
        for _, p in (st.session_state.projects or {}).items():
            if not isinstance(p, dict):
                continue
            ts = p.get("created_ts") or now_ts()
            title = p.get("title", "Untitled")
            p.setdefault("story_bible_id", hashlib.md5(f"sb|{title}|{ts}".encode("utf-8")).hexdigest()[:12])
            p.setdefault("story_bible_created_ts", ts)
            p.setdefault("voices", default_voice_vault())
            p.setdefault("style_banks", default_style_banks())
            if "story_bible_fingerprint" not in p:
                try:
                    p["story_bible_fingerprint"] = _fingerprint_story_bible(p.get("story_bible", {}) or {})
                except Exception:
                    p["story_bible_fingerprint"] = ""

        ab = payload.get("active_bay", "NEW")
        if ab not in BAYS:
            ab = "NEW"
        st.session_state.active_bay = ab

        ensure_bay_has_active_project(ab)
        pid = st.session_state.active_project_by_bay.get(ab)
        if pid:
            load_project_into_session(pid)
        else:
            if ab == "NEW":
                load_workspace_into_session()
                st.session_state.voice_status = "NEW: (Story Bible workspace)"
            else:
                switch_bay(ab)

        saved_at = (payload.get("meta", {}) or {}).get("saved_at", "")
        src = "autosave" if loaded_from == "primary" else "backup autosave"
        st.session_state.voice_status = f"Loaded {src} ({saved_at})."
        st.session_state.last_saved_digest = _digest(_payload())
    except Exception as e:
        st.session_state.voice_status = f"Load warning: {e}"
        _boot_new()


if "did_load_autosave" not in st.session_state:
    st.session_state.did_load_autosave = True
    try:
        load_all_from_disk()
    except Exception as e:
        st.error(f"❌ Failed to load autosave: {str(e)}")
        logger.error(f"Autosave load error: {e}\n{traceback.format_exc()}")
        st.write("Initializing fresh workspace...")
        try:
            _boot_new()
        except Exception as e2:
            st.error(f"❌ Failed to initialize workspace: {str(e2)}")
            st.stop()


def autosave() -> None:
    st.session_state.autosave_time = datetime.now().strftime("%H:%M:%S")
    save_all_to_disk()


# ============================================================
# IMPORT / EXPORT
# ============================================================
def _read_uploaded_text(uploaded) -> Tuple[str, str]:
    """Read .txt/.md/.docx from Streamlit UploadedFile."""
    if uploaded is None:
        return "", ""
    name = getattr(uploaded, "name", "") or ""
    raw = uploaded.getvalue()
    if raw is None:
        return "", name
    if len(raw) > MAX_UPLOAD_BYTES:
        return "", name
    ext = os.path.splitext(name.lower())[1]

    if ext in (".txt", ".md", ".markdown", ".text", ""):
        try:
            return raw.decode("utf-8"), name
        except Exception:
            return raw.decode("utf-8", errors="ignore"), name

    if ext == ".docx":
        try:
            from docx import Document  # python-docx

            doc = Document(BytesIO(raw))
            parts = []
            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t:
                    parts.append(t)
            return "\n\n".join(parts), name
        except Exception:
            try:
                return raw.decode("utf-8", errors="ignore"), name
            except Exception:
                return "", name

    try:
        return raw.decode("utf-8", errors="ignore"), name
    except Exception:
        return "", name


def _sb_sections_from_text_heuristic(text: str) -> Dict[str, str]:
    t = _normalize_text(text)
    if not t:
        return {"synopsis": "", "genre_style_notes": "", "world": "", "characters": "", "outline": ""}

    heading_map = {
        "synopsis": ["synopsis", "premise", "logline"],
        "genre_style_notes": ["genre", "style", "tone", "voice"],
        "world": ["world", "setting", "lore"],
        "characters": ["characters", "cast"],
        "outline": ["outline", "beats", "plot", "structure"],
    }

    lines = t.splitlines()
    buckets = {k: [] for k in heading_map.keys()}
    current = None

    def _match_heading(line: str) -> Optional[str]:
        l = re.sub(r"^[#\-\*\s]+", "", (line or "").strip()).lower()
        l = re.sub(r"[:\-\s]+$", "", l)
        for key, aliases in heading_map.items():
            if any(l == a or l.startswith(a + " ") for a in aliases):
                return key
        return None

    for line in lines:
        key = _match_heading(line)
        if key:
            current = key
            continue
        if current:
            buckets[current].append(line)
        else:
            buckets["synopsis"].append(line)

    return {k: _normalize_text("\n".join(v)) for k, v in buckets.items()}


def _extract_json_object(s: str) -> Optional[Dict[str, Any]]:
    if not s:
        return None
    s2 = re.sub(r"^```(?:json)?\s*|\s*```$", "", s.strip(), flags=re.IGNORECASE | re.MULTILINE)
    m = re.search(r"\{.*\}", s2, flags=re.DOTALL)
    if not m:
        return None
    try:
        obj = json.loads(m.group(0))
        return obj if isinstance(obj, dict) else None
    except Exception:
        return None


def story_bible_markdown(title: str, sb: Dict[str, str], meta: Dict[str, Any]) -> str:
    t = title or "Untitled"

    def sec(h: str, k: str) -> str:
        body = (sb.get(k, "") or "").strip()
        return f"## {h}\n\n{body}\n" if body else f"## {h}\n\n\n"

    front = f"# Story Bible — {t}\n\n- Exported: {now_ts()}\n"
    for mk, mv in (meta or {}).items():
        front += f"- {mk}: {mv}\n"
    front += "\n"
    return front + "\n".join(
        [
            sec("Synopsis", "synopsis"),
            sec("Genre / Style Notes", "genre_style_notes"),
            sec("World", "world"),
            sec("Characters", "characters"),
            sec("Outline", "outline"),
        ]
    )


def draft_markdown(title: str, draft: str, meta: Dict[str, Any]) -> str:
    t = title or "Untitled"
    front = f"# Draft — {t}\n\n- Exported: {now_ts()}\n"
    for mk, mv in (meta or {}).items():
        front += f"- {mk}: {mv}\n"
    front += "\n---\n\n"
    return front + (draft or "")


def make_project_bundle(pid: str) -> Dict[str, Any]:
    p = (st.session_state.projects or {}).get(pid, {}) or {}
    return {"meta": {"exported_at": now_ts(), "type": "project_bundle", "version": "1"}, "project": p}


def make_library_bundle() -> Dict[str, Any]:
    if in_workspace_mode():
        save_workspace_from_session()
    else:
        save_session_into_project()
    return {
        "meta": {"exported_at": now_ts(), "type": "library_bundle", "version": "1"},
        "active_bay": st.session_state.active_bay,
        "active_project_by_bay": st.session_state.active_project_by_bay,
        "sb_workspace": st.session_state.sb_workspace,
        "projects": st.session_state.projects,
    }


def _new_pid_like(seed: str) -> str:
    return hashlib.md5(f"{seed}|{now_ts()}".encode("utf-8")).hexdigest()[:12]


def import_project_bundle(bundle: Dict[str, Any], target_bay: str = "NEW", rename: str = "") -> Optional[str]:
    if not isinstance(bundle, dict):
        return None
    proj = bundle.get("project")
    if not isinstance(proj, dict):
        return None

    pid = str(proj.get("id") or _new_pid_like("import"))
    if pid in (st.session_state.projects or {}):
        pid = _new_pid_like(pid)

    proj = json.loads(json.dumps(proj))
    proj["id"] = pid
    if rename.strip():
        proj["title"] = rename.strip()
    if target_bay in BAYS:
        proj["bay"] = target_bay
    proj["updated_ts"] = now_ts()

    ts = proj.get("created_ts") or now_ts()
    title = proj.get("title", "Untitled")
    proj.setdefault("story_bible_id", hashlib.md5(f"sb|{title}|{ts}".encode("utf-8")).hexdigest()[:12])
    proj.setdefault("story_bible_created_ts", ts)
    # story_bible_binding and locks removed: Story Bible is always editable
    proj.setdefault("voices", default_voice_vault())
    proj.setdefault("style_banks", default_style_banks())
    proj.setdefault("story_bible_fingerprint", "")

    st.session_state.projects[pid] = proj
    st.session_state.active_project_by_bay[proj.get("bay", "NEW")] = pid
    return pid


def import_library_bundle(bundle: Dict[str, Any]) -> int:
    if not isinstance(bundle, dict):
        return 0
    projs = bundle.get("projects")
    if not isinstance(projs, dict):
        return 0
    imported = 0
    for _, proj in projs.items():
        if not isinstance(proj, dict):
            continue
        pid = import_project_bundle({"project": proj}, target_bay=proj.get("bay", "NEW"), rename="")
        if pid:
            imported += 1

    w = bundle.get("sb_workspace")
    if isinstance(w, dict) and w.get("workspace_story_bible_id"):
        cur = st.session_state.sb_workspace or default_story_bible_workspace()
        cur_sb = (cur.get("story_bible", {}) or {})
        cur_empty = not any((cur_sb.get(k, "") or "").strip() for k in ["synopsis", "genre_style_notes", "world", "characters", "outline"])
        if cur_empty:
            st.session_state.sb_workspace = w
    return imported


# ============================================================
# JUNK DRAWER COMMANDS
# ============================================================
CMD_FIND = re.compile(r"^\s*/find\s*:\s*(.+)$", re.IGNORECASE)
CMD_CREATE = re.compile(r"^\s*/create\s*:\s*(.+)$", re.IGNORECASE)
CMD_PROMOTE = re.compile(r"^\s*/promote\s*$", re.IGNORECASE)


def _run_find(term: str) -> str:
    term = (term or "").strip()
    if not term:
        return "Find: missing search term. Use /find: word"

    def _hits(label: str, text: str) -> List[str]:
        lines = (text or "").splitlines()
        out = []
        for i, line in enumerate(lines, start=1):
            if term.lower() in (line or "").lower():
                out.append(f"{label} L{i}: {line.strip()}")
            if len(out) >= 20:
                break
        return out

    hits = []
    hits += _hits("DRAFT", st.session_state.main_text)
    hits += _hits("SYNOPSIS", st.session_state.synopsis)
    hits += _hits("WORLD", st.session_state.world)
    hits += _hits("CHARS", st.session_state.characters)
    hits += _hits("OUTLINE", st.session_state.outline)

    if not hits:
        return f"Find: no matches for '{term}'."
    return "\n".join(hits[:30])


def handle_junk_commands() -> None:
    raw = (st.session_state.junk or "").strip()
    if not raw:
        return

    m = CMD_CREATE.match(raw)
    if m:
        title = m.group(1).strip()
        pid = create_project_from_current_bible(title)
        st.session_state.active_project_by_bay["NEW"] = pid
        st.session_state.active_bay = "NEW"
        load_project_into_session(pid)
        st.session_state.voice_status = f"Created project in NEW: {st.session_state.project_title}"
        st.session_state.last_action = "Create Project"
        st.session_state.junk = ""
        autosave()
        return

    if CMD_PROMOTE.match(raw):
        pid = st.session_state.project_id
        bay = st.session_state.active_bay
        nb = next_bay(bay)
        if not pid or not nb:
            st.session_state.tool_output = "Promote: no project selected, or already FINAL."
            st.session_state.voice_status = "Promote blocked."
            st.session_state.junk = ""
            autosave()
            return
        save_session_into_project()
        promote_project(pid, nb)
        st.session_state.active_project_by_bay[nb] = pid
        switch_bay(nb)
        st.session_state.voice_status = f"Promoted → {nb}: {st.session_state.project_title}"
        st.session_state.last_action = f"Promote → {nb}"
        st.session_state.junk = ""
        autosave()
        return

    m = CMD_FIND.match(raw)
    if m:
        st.session_state.tool_output = _clamp_text(_run_find(m.group(1)))
        st.session_state.voice_status = "Find complete"
        st.session_state.last_action = "Find"
        st.session_state.junk = ""
        autosave()
        return


# Run junk commands early (before widgets instantiate)
handle_junk_commands()


# ============================================================
# AI BRIEF + CALL
# ============================================================
def _story_bible_text() -> str:
    sb = []
    if (st.session_state.synopsis or "").strip():
        sb.append(f"SYNOPSIS:\n{st.session_state.synopsis.strip()}")
    if (st.session_state.genre_style_notes or "").strip():
        sb.append(f"GENRE/STYLE NOTES:\n{st.session_state.genre_style_notes.strip()}")
    if (st.session_state.world or "").strip():
        sb.append(f"WORLD:\n{st.session_state.world.strip()}")
    if (st.session_state.characters or "").strip():
        sb.append(f"CHARACTERS:\n{st.session_state.characters.strip()}")
    if (st.session_state.outline or "").strip():
        sb.append(f"OUTLINE:\n{st.session_state.outline.strip()}")
    return "\n\n".join(sb).strip() if sb else "— None provided —"


def build_partner_brief(action_name: str, lane: str) -> str:
    """
    ═══════════════════════════════════════════════════════════════
    CORE INTEGRATION HUB - Assembles all Voice Bible controls into
    a unified AI prompt. Used by ALL AI generation functions.
    ═══════════════════════════════════════════════════════════════
    """
    story_bible = _story_bible_text()
    vb = []
    if st.session_state.vb_style_on:
        vb.append(f"Writing Style: {st.session_state.writing_style} (intensity {st.session_state.style_intensity:.2f})")
    if st.session_state.vb_genre_on:
        vb.append(f"Genre Influence: {st.session_state.genre} (intensity {st.session_state.genre_intensity:.2f})")
    if st.session_state.vb_trained_on and st.session_state.trained_voice and st.session_state.trained_voice != "— None —":
        vb.append(f"Trained Voice: {st.session_state.trained_voice} (intensity {st.session_state.trained_intensity:.2f})")
    if st.session_state.vb_match_on and (st.session_state.voice_sample or "").strip():
        vb.append(f"Match Sample (intensity {st.session_state.match_intensity:.2f}):\n{st.session_state.voice_sample.strip()}")
    if st.session_state.vb_lock_on and (st.session_state.voice_lock_prompt or "").strip():
        vb.append(f"VOICE LOCK (strength {st.session_state.lock_intensity:.2f}):\n{st.session_state.voice_lock_prompt.strip()}")
    if st.session_state.vb_technical_on:
        vb.append(f"Technical Controls: POV={st.session_state.pov}, Tense={st.session_state.tense} (enforcement {st.session_state.technical_intensity:.2f})")
    voice_controls = "\n\n".join(vb).strip() if vb else "— None enabled —"

    # Engine Style (trainable banks) → Semantic retrieval of trained samples
    style_name = (st.session_state.writing_style or "").strip().upper()
    style_directive = ""
    style_exemplars: List[str] = []
    if st.session_state.vb_style_on and style_name in ENGINE_STYLES:
        style_directive = engine_style_directive(style_name, float(st.session_state.style_intensity), lane)
        ctx2 = (st.session_state.main_text or "")[-2500:]
        q2 = ctx2 if ctx2.strip() else (st.session_state.synopsis or "")
        k = 1 + int(max(0.0, min(1.0, float(st.session_state.style_intensity))) * 2.0)
        style_exemplars = retrieve_style_exemplars(style_name, lane, q2, k=k)  # ← INTELLIGENT RETRIEVAL

    # Trained Voice → Vector-based semantic matching
    exemplars: List[str] = []
    tv = st.session_state.trained_voice
    if st.session_state.vb_trained_on and tv and tv != "— None —":
        ctx = (st.session_state.main_text or "")[-2500:]
        query = ctx if ctx.strip() else st.session_state.synopsis
        exemplars = retrieve_mixed_exemplars(tv, lane, query)  # ← ADAPTIVE RETRIEVAL
    ex_block = "\n\n---\n\n".join(exemplars) if exemplars else "— None —"
    style_ex_block = "\n\n---\n\n".join(style_exemplars) if style_exemplars else "— None —"

    ai_x = float(st.session_state.ai_intensity)
    return f"""
YOU ARE OLIVETTI: the author's personal writing and editing partner.
Professional output only. No UI talk. No process talk.

STORY BIBLE IS CANON + IDEA BANK.
When generating NEW material, pull at least 2 concrete specifics from the Story Bible.
Never contradict canon. Never add random characters/places unless Story Bible supports it.

LANE: {lane}

AI INTENSITY: {ai_x:.2f}
INTENSITY PROFILE: {intensity_profile(ai_x)}

VOICE CONTROLS:
{voice_controls}

ENGINE STYLE DIRECTIVE:
{style_directive if style_directive else "— None —"}

STYLE EXEMPLARS (mimic cadence/diction, not content):
{style_ex_block}

TRAINED EXEMPLARS (mimic patterns, not content):
{ex_block}

STORY BIBLE:
{story_bible}

ACTION: {action_name}
""".strip()


def call_openai(system_brief: str, user_task: str, text: str) -> str:
    """
    ═══════════════════════════════════════════════════════════════
    UNIFIED AI GATEWAY - Single entry point for ALL AI generation.
    Applies AI Intensity → Temperature conversion automatically.
    ═══════════════════════════════════════════════════════════════
    """
    key = require_openai_key()
    try:
        from openai import OpenAI
    except Exception as e:
        raise RuntimeError("OpenAI SDK not installed. Add to requirements.txt: openai") from e

    try:
        client = OpenAI(api_key=key, timeout=60)
    except TypeError:
        client = OpenAI(api_key=key)

    resp = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[
            {"role": "system", "content": system_brief},
            {"role": "user", "content": f"{user_task}\n\nDRAFT:\n{text.strip()}"},
        ],
        temperature=temperature_from_intensity(st.session_state.ai_intensity),  # ← AI INTENSITY → TEMPERATURE
    )
    result = (resp.choices[0].message.content or "").strip()
    logger.info(f"call_openai returned {len(result)} chars: {result[:100] if result else 'EMPTY'}")
    return result


def local_cleanup(text: str) -> str:
    t = (text or "")
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    t = re.sub(r"[ \t]+\n", "\n", t)
    t = re.sub(r"\n{4,}", "\n\n\n", t)
    t = re.sub(r"\s+([,.;:!?])", r"\1", t)
    t = re.sub(r"([,.;:!?])([A-Za-z0-9])", r"\1 \2", t)
    t = re.sub(r"\.\.\.", "…", t)
    t = re.sub(r"\s*--\s*", " — ", t)
    t = re.sub(r"[ \t]{2,}", " ", t)
    return t.strip()


def sb_breakdown_ai(text: str) -> Dict[str, str]:
    prompt = """Break the provided source document into these Story Bible sections.
Return STRICT JSON with these exact keys:
synopsis, genre_style_notes, world, characters, outline

Rules:
- Use clear, concrete specifics (names, places, stakes).
- Preserve existing proper nouns from the text.
- If a section is unknown, return an empty string for it.
- No commentary. JSON only.
"""
    try:
        out = call_openai(
            system_brief="You are a precise literary analyst. You output only valid JSON.",
            user_task=prompt,
            text=text,
        )
        obj = _extract_json_object(out) or {}
        return {
            "synopsis": _normalize_text(str(obj.get("synopsis", ""))),
            "genre_style_notes": _normalize_text(str(obj.get("genre_style_notes", ""))),
            "world": _normalize_text(str(obj.get("world", ""))),
            "characters": _normalize_text(str(obj.get("characters", ""))),
            "outline": _normalize_text(str(obj.get("outline", ""))),
        }
    except Exception:
        return _sb_sections_from_text_heuristic(text)


def _merge_section(existing: str, incoming: str, mode: str) -> str:
    ex = (existing or "").strip()
    inc = (incoming or "").strip()
    if mode == "Replace":
        return inc
    if not ex:
        return inc
    if not inc:
        return ex
    return (ex + "\n\n" + inc).strip()


def format_manuscript_standard(title: str, author: str, text: str, word_count: int) -> str:
    """
    Format text according to industry-standard manuscript guidelines:
    - Title page with word count
    - Proper spacing and indentation
    - Chapter breaks
    - Page headers
    """
    lines = []
    
    # Title Page
    lines.append(f"{author}")
    lines.append(f"{word_count} words")
    lines.append("")
    lines.append("")
    lines.append("")
    lines.append("")
    lines.append(f"{title.upper()}")
    lines.append("")
    lines.append(f"by {author}")
    lines.append("")
    lines.append("")
    lines.append("\n" * 10)  # Page break
    lines.append("")
    
    # Format body text
    paragraphs = text.split('\n\n')
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        # Detect chapter headings (all caps or starts with "Chapter")
        if para.isupper() or para.startswith("Chapter") or para.startswith("CHAPTER"):
            lines.append("\n" * 3)  # Extra spacing before chapter
            lines.append(para)
            lines.append("")
        else:
            # Regular paragraph with proper indentation
            lines.append(f"    {para}")
            lines.append("")
    
    return "\n".join(lines)


def format_ebook_html(title: str, author: str, text: str) -> str:
    """
    Format text as clean HTML suitable for ebook conversion.
    Includes proper semantic markup for chapters, paragraphs, etc.
    """
    html = []
    html.append('<!DOCTYPE html>')
    html.append('<html lang="en">')
    html.append('<head>')
    html.append('    <meta charset="UTF-8">')
    html.append('    <meta name="viewport" content="width=device-width, initial-scale=1.0">')
    html.append(f'    <title>{title}</title>')
    html.append('    <style>')
    html.append('        body { font-family: Georgia, serif; line-height: 1.6; max-width: 600px; margin: 0 auto; padding: 20px; }')
    html.append('        h1 { text-align: center; margin: 2em 0; font-size: 2em; }')
    html.append('        h2 { text-align: center; margin: 2em 0 1em 0; page-break-before: always; }')
    html.append('        p { text-indent: 2em; margin: 0 0 1em 0; }')
    html.append('        p.first { text-indent: 0; }')
    html.append('        .title-page { text-align: center; margin: 4em 0; page-break-after: always; }')
    html.append('        .author { font-size: 1.2em; margin-top: 1em; }')
    html.append('    </style>')
    html.append('</head>')
    html.append('<body>')
    
    # Title page
    html.append('    <div class="title-page">')
    html.append(f'        <h1>{title}</h1>')
    html.append(f'        <p class="author">by {author}</p>')
    html.append('    </div>')
    
    # Body text
    paragraphs = text.split('\n\n')
    in_chapter = False
    first_in_chapter = True
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        # Detect chapter headings
        if para.isupper() or para.startswith("Chapter") or para.startswith("CHAPTER"):
            html.append(f'    <h2>{para}</h2>')
            in_chapter = True
            first_in_chapter = True
        else:
            # Regular paragraph
            p_class = ' class="first"' if first_in_chapter else ''
            html.append(f'    <p{p_class}>{para}</p>')
            first_in_chapter = False
    
    html.append('</body>')
    html.append('</html>')
    
    return '\n'.join(html)


# ============================================================
# STORY BIBLE AI GENERATION
# ============================================================
def generate_story_bible_section(section_type: str) -> None:
    """Generate content for a specific Story Bible section using AI.
    RESPECTS ALL VOICE BIBLE SETTINGS - uses same engine as writing actions."""
    if not has_openai_key():
        st.session_state.tool_output = f"AI generation requires OPENAI_API_KEY to be configured."
        st.session_state.voice_status = f"{section_type}: AI unavailable"
        autosave()
        return

    # Gather context from existing Story Bible sections and draft
    context_parts = []
    if st.session_state.main_text:
        context_parts.append(f"DRAFT TEXT:\n{st.session_state.main_text[:3000]}")
    if st.session_state.synopsis and section_type != "Synopsis":
        context_parts.append(f"SYNOPSIS:\n{st.session_state.synopsis}")
    if st.session_state.genre_style_notes and section_type != "Genre/Style":
        context_parts.append(f"GENRE/STYLE:\n{st.session_state.genre_style_notes}")
    if st.session_state.world and section_type != "World":
        context_parts.append(f"WORLD:\n{st.session_state.world}")
    if st.session_state.characters and section_type != "Characters":
        context_parts.append(f"CHARACTERS:\n{st.session_state.characters}")
    if st.session_state.outline and section_type != "Outline":
        context_parts.append(f"OUTLINE:\n{st.session_state.outline}")
    
    context = "\n\n---\n\n".join(context_parts) if context_parts else "No context available yet."

    # Section-specific generation prompts
    prompts = {
        "Synopsis": "Write a concise, compelling 2-3 paragraph synopsis that captures the story's core conflict, main characters, and stakes. Be specific with names and situations.",
        "Genre/Style": "Describe the genre, tone, voice, and stylistic approach for this story. Include specific markers like 'hardboiled detective', 'lyrical literary fiction', 'tight thriller prose', etc.",
        "World": "Detail the world, setting, and key locations. Include rules, atmosphere, time period, and any special systems (magic, technology, social structures). Be concrete and specific.",
        "Characters": "List and describe the main characters with names, roles, relationships, motivations, and key traits. Format as a character list with details.",
        "Outline": "Create a story outline with acts, major beats, key scenes, and turning points. Structure it clearly with progression from beginning to end.",
    }

    try:
        task = prompts.get(section_type, f"Generate {section_type} content for the Story Bible.")
        
        # BUILD FULL VOICE BIBLE BRIEF - same as writing actions
        # This ensures Story Bible generation respects ALL Voice Bible controls
        vb_controls = []
        if st.session_state.vb_style_on:
            vb_controls.append(f"Writing Style: {st.session_state.writing_style} (intensity {st.session_state.style_intensity:.2f})")
        if st.session_state.vb_genre_on:
            vb_controls.append(f"Genre Influence: {st.session_state.genre} (intensity {st.session_state.genre_intensity:.2f})")
        if st.session_state.vb_trained_on and st.session_state.trained_voice and st.session_state.trained_voice != "— None —":
            vb_controls.append(f"Trained Voice: {st.session_state.trained_voice} (intensity {st.session_state.trained_intensity:.2f})")
        if st.session_state.vb_match_on and (st.session_state.voice_sample or "").strip():
            vb_controls.append(f"Match Sample (intensity {st.session_state.match_intensity:.2f}):\n{st.session_state.voice_sample.strip()}")
        if st.session_state.vb_lock_on and (st.session_state.voice_lock_prompt or "").strip():
            vb_controls.append(f"VOICE LOCK (strength {st.session_state.lock_intensity:.2f}):\n{st.session_state.voice_lock_prompt.strip()}")
        
        voice_brief = "\n\n".join(vb_controls) if vb_controls else "— No Voice Bible controls active —"
        
        ai_x = float(st.session_state.ai_intensity)
        brief = f"""You are a story development expert helping build a comprehensive Story Bible.

AI INTENSITY: {ai_x:.2f}
INTENSITY PROFILE: {intensity_profile(ai_x)}

VOICE CONTROLS (apply these to your output):
{voice_brief}

POV: {st.session_state.pov}
TENSE: {st.session_state.tense}

EXISTING CONTEXT:
{context}"""
        
        result = call_openai(brief, task, "")
        
        if result:
            if section_type == "Synopsis":
                st.session_state.synopsis = result.strip()
            elif section_type == "Genre/Style":
                st.session_state.genre_style_notes = result.strip()
            elif section_type == "World":
                st.session_state.world = result.strip()
            elif section_type == "Characters":
                st.session_state.characters = result.strip()
            elif section_type == "Outline":
                st.session_state.outline = result.strip()
            
            st.session_state.voice_status = f"Generated: {section_type} (Voice Bible applied)"
            st.session_state.last_action = f"Generate {section_type}"
            autosave()
        else:
            st.session_state.voice_status = f"AI generation failed: {section_type}"
    except Exception as e:
        st.session_state.tool_output = f"AI generation error: {str(e)}"
        st.session_state.voice_status = f"{section_type}: generation failed"
        autosave()


# ============================================================
# ACTIONS (queued for Streamlit safety)
# ============================================================
def partner_action(action: str) -> None:
    """
    ═══════════════════════════════════════════════════════════════
    WRITING ACTIONS ROUTER - All bottom bar functions route here.
    Uses build_partner_brief() → Ensures ALL actions use Voice Bible.
    ═══════════════════════════════════════════════════════════════
    """
    logger.info(f"partner_action called: {action}")
    
    # Use selection text if provided, otherwise use full draft
    selection = (st.session_state.get("selection_text") or "").strip()
    text = selection if selection else (st.session_state.main_text or "")
    is_selection = bool(selection)
    
    lane = current_lane_from_draft(text)
    brief = build_partner_brief(action, lane=lane)  # ← UNIFIED BRIEF with ALL Voice Bible controls
    use_ai = has_openai_key()
    logger.info(f"use_ai = {use_ai}, text length = {len(text)}, is_selection = {is_selection}")

    def show_preview(result: str, action_name: str) -> None:
        """Show AI output in preview window for user review"""
        logger.info(f"show_preview called: action={action_name}, result_length={len(result) if result else 0}, result_preview={result[:100] if result else 'None'}")
        if result and result.strip():
            # Write AI output directly to main_text (no preview)
            if action_name in ["Write"]:
                if (st.session_state.main_text or "").strip():
                    st.session_state.main_text = (st.session_state.main_text.rstrip() + "\n\n" + result.strip()).strip()
                else:
                    st.session_state.main_text = result.strip()
            else:
                # For selection-based edits, just show in preview - user copies manually
                # For full draft edits, replace
                if not st.session_state.selection_text.strip():
                    st.session_state.main_text = result.strip()
                else:
                    st.session_state.main_text = st.session_state.main_text

            st.session_state.voice_status = f"{action_name} complete and added to draft (AI output)"
            st.session_state.last_action = f"{action_name} (auto-accept)"
            autosave()
            logger.info(f"show_preview: main_text updated, length = {len(st.session_state.main_text)}")
            st.session_state.tool_output = f"✓ {action_name} complete. Output added to draft."
            import streamlit as stlib
            stlib.rerun()
        else:
            logger.warning(f"show_preview: empty or None result for {action_name}")
            st.session_state.tool_output = f"⚠ {action_name} returned empty result."

    def apply_replace(result: str) -> None:
        if result and result.strip():
            st.session_state.main_text = result.strip()
            st.session_state.last_action = action
            # Show Voice Bible was applied
            vb_summary = _get_voice_bible_summary()
            st.session_state.voice_status = f"{action} complete ({vb_summary})"
            autosave()

    def apply_append(result: str) -> None:
        if result and result.strip():
            if (st.session_state.main_text or "").strip():
                st.session_state.main_text = (st.session_state.main_text.rstrip() + "\n\n" + result.strip()).strip()
            else:
                st.session_state.main_text = result.strip()
            st.session_state.last_action = action
            # Show Voice Bible was applied
            vb_summary = _get_voice_bible_summary()
            st.session_state.voice_status = f"{action} complete ({vb_summary})"
            autosave()

    try:
        if action == "Write":
            if use_ai:
                task = (
                    f"Continue decisively in lane ({lane}). Add 1–3 paragraphs that advance the scene. "
                    "MANDATORY: incorporate at least 2 Story Bible specifics. "
                    "No recap. No planning. Just prose."
                )
                out = call_openai(brief, task, text if text.strip() else "Start the opening.")
                show_preview(out, "Write")
            else:
                st.session_state.tool_output = "Write requires OPENAI_API_KEY to be configured."
                st.session_state.voice_status = "Write: API key missing"
                autosave()
            return

        if action == "Rewrite":
            if use_ai:
                task = f"Rewrite for professional quality in lane ({lane}). Preserve meaning and canon. Return full revised text."
                out = call_openai(brief, task, text)
                show_preview(out, "Rewrite")
            else:
                show_preview(local_cleanup(text), "Rewrite")
            return

        if action == "Expand":
            if use_ai:
                task = f"Expand with meaningful depth in lane ({lane}). No padding. Preserve canon. Return full revised text."
                out = call_openai(brief, task, text)
                show_preview(out, "Expand")
            else:
                st.session_state.tool_output = "Expand requires OPENAI_API_KEY to be configured."
                st.session_state.voice_status = "Expand: API key missing"
                autosave()
            return

        if action == "Rephrase":
            if use_ai:
                task = f"Replace the final sentence with a stronger one (same meaning) in lane ({lane}). Return full text."
                out = call_openai(brief, task, text)
                show_preview(out, "Rephrase")
            else:
                st.session_state.tool_output = "Rephrase requires OPENAI_API_KEY to be configured."
                st.session_state.voice_status = "Rephrase: API key missing"
                autosave()
            return

        if action == "Describe":
            if use_ai:
                task = f"Add vivid controlled description in lane ({lane}). Preserve pace and canon. Return full revised text."
                out = call_openai(brief, task, text)
                show_preview(out, "Describe")
            else:
                st.session_state.tool_output = "Describe requires OPENAI_API_KEY to be configured."
                st.session_state.voice_status = "Describe: API key missing"
                autosave()
            return

        if action in ("Spell", "Grammar"):
            cleaned = local_cleanup(text)
            if use_ai:
                task = "Copyedit spelling/grammar/punctuation. Preserve voice. Return full revised text."
                out = call_openai(brief, task, cleaned)
                show_preview(out if out else cleaned, action)
            else:
                show_preview(cleaned, action)
            return

        if action == "Synonym":
            # Tool output only (does not change draft)
            last = ""
            m = re.search(r"([A-Za-z']{3,})\W*$", text.strip())
            if m:
                last = m.group(1)
            if not last:
                st.session_state.tool_output = "Synonym: couldn't detect a target word (try ending with a word)."
                st.session_state.voice_status = "Synonym: no target"
                autosave()
                return
            if use_ai:
                task = (
                    f"Give 12 strong synonyms for '{last}'. "
                    "Group them by nuance (formal, punchy, poetic, archaic, etc). "
                    "No filler." 
                )
                out = call_openai(brief, task, text)
                st.session_state.tool_output = _clamp_text(out)
            else:
                st.session_state.tool_output = f"Synonym requires OPENAI_API_KEY (target word: {last})."
            st.session_state.voice_status = f"Synonym: {last}"
            st.session_state.last_action = "Synonym"
            autosave()
            return

        if action == "Sentence":
            # Tool output only
            last_sentence = ""
            sentences = re.split(r"(?<=[.!?])\s+", text.strip())
            if sentences:
                last_sentence = sentences[-1].strip()
            if not last_sentence:
                st.session_state.tool_output = "Sentence: couldn't detect a final sentence."
                st.session_state.voice_status = "Sentence: no target"
                autosave()
                return
            if use_ai:
                task = (
                    "Provide 8 alternative rewrites of the final sentence. "
                    "Keep meaning. Vary rhythm and diction. Return as a numbered list."
                )
                out = call_openai(brief, task, text)
                st.session_state.tool_output = _clamp_text(out)
            else:
                st.session_state.tool_output = "Sentence requires OPENAI_API_KEY."
            st.session_state.voice_status = "Sentence options"
            st.session_state.last_action = "Sentence"
            autosave()
            return

    except Exception as e:
        import traceback
        msg = str(e)
        logger.error(f"partner_action error: {msg}\n{traceback.format_exc()}")
        
        if ("insufficient_quota" in msg) or ("exceeded your current quota" in msg.lower()):
            st.session_state.voice_status = "Engine: OpenAI quota exceeded."
            st.session_state.tool_output = _clamp_text(
                "OpenAI returned a quota/billing error.\n\nFix:\n• Confirm your API key is correct\n• Check billing/usage limits\n• Or swap to a different key in Streamlit Secrets"
            )
        elif "OPENAI_API_KEY not set" in msg:
            st.session_state.voice_status = "Engine: missing OPENAI_API_KEY."
            st.session_state.tool_output = "Set OPENAI_API_KEY in Streamlit Secrets (or environment) to enable AI."
        elif "401" in msg or "Unauthorized" in msg:
            st.session_state.voice_status = "Engine: API key rejected (401)"
            st.session_state.tool_output = _clamp_text(
                f"OpenAI rejected the API key (401 Unauthorized).\n\nError: {msg}\n\nFix:\n• Check your API key at platform.openai.com\n• Make sure it's active and has credits\n• Key should be in .streamlit/secrets.toml"
            )
        else:
            st.session_state.voice_status = f"Engine: {msg[:50]}"
            st.session_state.tool_output = _clamp_text(f"ERROR:\n{msg}\n\nFull trace:\n{traceback.format_exc()}")
        autosave()


def queue_action(action: str) -> None:
    st.session_state.pending_action = action


def run_pending_action() -> None:
    action = st.session_state.get("pending_action")
    if not action:
        return
    st.session_state.pending_action = None

    # Non-engine UI hint (keeps session_state mutations pre-widget)
    if action == "__FIND_HINT__":
        st.session_state.tool_output = "Find is wired via /find: in Junk Drawer."
        st.session_state.voice_status = "Find: use /find: in Junk Drawer."
        st.session_state.last_action = "Find (hint)"
        autosave()
        return

    # __VAULT_CLEAR_SAMPLE__ logic removed: sample box is not cleared after add, only on Accept if desired.
    logger.info(f"run_pending_action: calling partner_action({action})")
    try:
        partner_action(action)
        # Force a rerun after actions that update preview to ensure UI shows the result
        preview_content = st.session_state.get("ai_preview", "")
        logger.info(f"run_pending_action: after partner_action, ai_preview length = {len(preview_content)}")
        if preview_content:
            logger.info(f"run_pending_action: triggering rerun to show preview")
            st.rerun()
    except Exception as e:
        logger.error(f"Error in partner_action({action}): {str(e)}", exc_info=True)
        st.session_state.tool_output = f"❌ Error: {str(e)}"
        st.session_state.voice_status = f"{action} failed"


# ============================================================
# UI — TOP BAR
# ============================================================
top = st.container()
with top:
    cols = st.columns([1, 1, 1, 1, 2])

    if cols[0].button("🆕 New", key="bay_new"):
        switch_bay("NEW")
        save_all_to_disk(force=True)

    if cols[1].button("✏️ Rough", key="bay_rough"):
        switch_bay("ROUGH")
        save_all_to_disk(force=True)

    if cols[2].button("🛠 Edit", key="bay_edit"):
        switch_bay("EDIT")
        save_all_to_disk(force=True)

    if cols[3].button("✅ Final", key="bay_final"):
        switch_bay("FINAL")
        save_all_to_disk(force=True)

    cols[4].markdown(
        f"""
        <div style='text-align:right;font-size:12px;'>
            Bay: <b>{st.session_state.active_bay}</b>
            &nbsp;•&nbsp; Project: <b>{st.session_state.project_title}</b>
            &nbsp;•&nbsp; Autosave: {st.session_state.autosave_time or '—'}
            <br/>AI Intensity: {float(st.session_state.ai_intensity):.2f}
            &nbsp;•&nbsp; Status: {st.session_state.voice_status}
        </div>
        """,
        unsafe_allow_html=True,
    )

st.divider()

# ============================================================
# LOCKED LAYOUT (same ratios)
# ============================================================
left, center, right = st.columns([1.2, 3.2, 1.6])


# ============================================================
# LEFT — STORY BIBLE
# ============================================================
with left:
    st.subheader("📖 Story Bible")

    bay = st.session_state.active_bay
    bay_projects = list_projects_in_bay(bay)

    # Build a stable selector with unique labels (handles duplicate titles)
    if bay == "NEW":
        items: List[Tuple[Optional[str], str]] = [(None, "— (Story Bible workspace) —")] + [(pid, title) for pid, title in bay_projects]
    else:
        items = [(None, "— (none) —")] + [(pid, title) for pid, title in bay_projects]

    # Disambiguate duplicate titles
    seen: Dict[str, int] = {}
    labels: List[str] = []
    ids: List[Optional[str]] = []
    for pid, title in items:
        base = title
        seen[base] = seen.get(base, 0) + 1
        label = base if seen[base] == 1 else f"{base}  ·  {str(pid)[-4:]}"
        labels.append(label)
        ids.append(pid)

    current_pid = st.session_state.project_id if (st.session_state.project_id in ids) else None
    current_idx = ids.index(current_pid) if current_pid in ids else 0

    sel = st.selectbox("Current Bay Project", labels, index=current_idx, key="bay_project_selector")
    sel_pid = ids[labels.index(sel)] if sel in labels else None

    if sel_pid:
        p = st.session_state.projects.get(sel_pid, {}) or {}
        sbid = p.get("story_bible_id", "—")
        sbts = p.get("story_bible_created_ts", "—")
        bind = p.get("story_bible_binding", {}) or {}
        src = bind.get("source", "—")
        st.caption(f"Locked Story Bible → Project • Bible ID: **{sbid}** • Created: **{sbts}** • Source: **{src}**")
    else:
        w = st.session_state.sb_workspace or default_story_bible_workspace()
        st.caption(
            f"Workspace Story Bible (not linked yet) • Bible ID: **{w.get('workspace_story_bible_id','—')}** • Created: **{w.get('workspace_story_bible_created_ts','—')}**"
        )

    # Switch context if changed
    if sel_pid != st.session_state.project_id:
        if in_workspace_mode():
            save_workspace_from_session()
        else:
            save_session_into_project()

        st.session_state.active_project_by_bay[bay] = sel_pid
        if sel_pid:
            load_project_into_session(sel_pid)
            st.session_state.voice_status = f"{bay}: {st.session_state.project_title}"
        else:
            st.session_state.project_id = None
            st.session_state.project_title = "—"
            if bay == "NEW":
                load_workspace_into_session()
                st.session_state.voice_status = "NEW: (Story Bible workspace)"
            else:
                st.session_state.main_text = ""
                st.session_state.synopsis = ""
                st.session_state.genre_style_notes = ""
                st.session_state.world = ""
                st.session_state.characters = ""
                st.session_state.outline = ""
                st.session_state.voice_sample = ""
                set_ai_intensity(0.75)
                st.session_state.voices = rebuild_vectors_in_voice_vault(default_voice_vault())
                st.session_state.voices_seeded = True
                st.session_state.voice_status = f"{bay}: (empty)"
        st.session_state.last_action = "Select Context"
        autosave()

    # Hard lock: Story Bible edits are locked per-project unless explicitly unlocked
    is_project = bool(st.session_state.project_id)
    sb_edit_unlocked = bool((st.session_state.locks or {}).get("sb_edit_unlocked", False))
    sb_locked = is_project and (not sb_edit_unlocked)

    with st.expander("🔒 Story Bible Hard Lock", expanded=False):
        if is_project:
            st.caption("Default is LOCKED. Unlock only when you intend to edit canon.")
            st.checkbox("Unlock Story Bible Editing", key="sb_unlock_cb")
            # sync the checkbox into locks
            st.session_state.locks["sb_edit_unlocked"] = bool(st.session_state.sb_unlock_cb)
            autosave()
        else:
            st.caption("Workspace is always editable. This lock applies after a project is created.")

    # ✅ AI Intensity - MASTER CONTROL for all AI generation
    st.divider()
    st.subheader("🎚️ AI Intensity (Master Control)")
    st.caption("⚙️ Controls ALL AI generation. Voice Bible settings apply on top of this base intensity.")
    
    ai_int_col1, ai_int_col2 = st.columns([3, 1])
    with ai_int_col1:
        st.slider(
            "AI Intensity",
            0.0,
            1.0,
            key="ai_intensity",
            help="0.0 = conservative/precise, 1.0 = bold/creative. Applies to EVERY AI action.",
            on_change=autosave,
        )
    with ai_int_col2:
        current_ai = float(st.session_state.ai_intensity)
        if current_ai <= 0.25:
            label, desc = "LOW", "Conservative"
        elif current_ai <= 0.60:
            label, desc = "MED", "Balanced"
        elif current_ai <= 0.85:
            label, desc = "HIGH", "Bold"
        else:
            label, desc = "MAX", "Creative"
        st.metric("Mode", label, desc)

    # Project controls
    action_cols = st.columns([1, 1])
    if bay == "NEW":
        label = "Start Project (Lock Bible → Project)" if in_workspace_mode() else "Create Project (from Bible)"
        if action_cols[0].button(label, key="create_project_btn"):
            title_guess = (st.session_state.synopsis.strip().splitlines()[0].strip() if st.session_state.synopsis.strip() else "New Project")
            pid = create_project_from_current_bible(title_guess)
            load_project_into_session(pid)
            st.session_state.voice_status = f"Created in NEW: {st.session_state.project_title}"
            st.session_state.last_action = "Create Project"
            autosave()
            st.rerun()

        if in_workspace_mode() and action_cols[1].button("New Story Bible (fresh ID)", key="new_workspace_bible_btn"):
            reset_workspace_story_bible(keep_templates=True)
            st.session_state.voice_status = "Workspace: new Story Bible minted"
            st.session_state.last_action = "New Story Bible"
            autosave()
            st.rerun()

        if action_cols[1].button("Promote → Rough", key="promote_new_to_rough"):
            if st.session_state.project_id:
                save_session_into_project()
                promote_project(st.session_state.project_id, "ROUGH")
                st.session_state.active_project_by_bay["ROUGH"] = st.session_state.project_id
                switch_bay("ROUGH")
                st.session_state.voice_status = f"Promoted → ROUGH: {st.session_state.project_title}"
                st.session_state.last_action = "Promote → ROUGH"
                autosave()
                st.rerun()
    elif bay in ("ROUGH", "EDIT"):
        nb = next_bay(bay)
        if action_cols[1].button(f"Promote → {nb.title()}", key=f"promote_{bay.lower()}"):
            if st.session_state.project_id and nb:
                save_session_into_project()
                promote_project(st.session_state.project_id, nb)
                st.session_state.active_project_by_bay[nb] = st.session_state.project_id
                switch_bay(nb)
                st.session_state.voice_status = f"Promoted → {nb}: {st.session_state.project_title}"
                st.session_state.last_action = f"Promote → {nb}"
                autosave()
                st.rerun()

    # Import / Export hub (restored)
    with st.expander("📦 Import / Export", expanded=False):
        tab_imp, tab_exp, tab_bundle = st.tabs(["Import", "Export", "Bundles"])

        with tab_imp:
            st.caption("Import a document into Draft or break it into Story Bible sections.")
            up = st.file_uploader("Upload (.txt, .md, .docx)", type=["txt", "md", "docx"], key="io_upload")
            paste = st.text_area("Paste text", key="io_paste", height=140)
            target = st.radio("Import target", ["Draft", "Story Bible"], horizontal=True, key="io_target")
            merge_mode = st.radio("Merge mode", ["Append", "Replace"], horizontal=True, key="io_merge")
            use_ai = st.checkbox(
                "Use AI Breakdown (Story Bible)",
                value=has_openai_key(),
                disabled=not has_openai_key(),
                help="Requires OPENAI_API_KEY. Falls back to heuristic if AI fails.",
                key="io_use_ai",
            )

            if st.button("Run Import", key="io_run_import"):
                src_file, name = _read_uploaded_text(up)
                src = _normalize_text(paste if (paste or "").strip() else src_file)
                if not src.strip():
                    st.session_state.tool_output = "Import: no text provided (or file too large)."
                    st.session_state.voice_status = "Import blocked"
                    autosave()
                elif target == "Draft":
                    if merge_mode == "Replace":
                        st.session_state.main_text = src
                    else:
                        st.session_state.main_text = (st.session_state.main_text.rstrip() + "\n\n" + src).strip() if (st.session_state.main_text or "").strip() else src
                    st.session_state.voice_status = f"Imported → Draft ({name or 'paste'})"
                    st.session_state.last_action = "Import → Draft"
                    autosave()
                else:
                    if sb_locked:
                        st.session_state.tool_output = "Story Bible is LOCKED. Unlock Story Bible Editing to import into it."
                        st.session_state.voice_status = "Import blocked (locked)"
                        autosave()
                    else:
                        sections = sb_breakdown_ai(src) if use_ai else _sb_sections_from_text_heuristic(src)
                        st.session_state.synopsis = _merge_section(st.session_state.synopsis, sections.get("synopsis", ""), merge_mode)
                        st.session_state.genre_style_notes = _merge_section(
                            st.session_state.genre_style_notes, sections.get("genre_style_notes", ""), merge_mode
                        )
                        st.session_state.world = _merge_section(st.session_state.world, sections.get("world", ""), merge_mode)
                        st.session_state.characters = _merge_section(st.session_state.characters, sections.get("characters", ""), merge_mode)
                        st.session_state.outline = _merge_section(st.session_state.outline, sections.get("outline", ""), merge_mode)
                        st.session_state.voice_status = f"Imported → Story Bible ({'AI' if use_ai else 'heuristic'})"
                        st.session_state.last_action = "Import → Story Bible"
                        autosave()

        with tab_exp:
            title = "Workspace" if in_workspace_mode() else st.session_state.project_title
            meta = {"Context": "Workspace" if in_workspace_mode() else "Project", "Bay": st.session_state.active_bay}
            if st.session_state.project_id:
                meta["Project ID"] = st.session_state.project_id
            stem = _safe_filename(title, "olivetti")

            # Ensure latest writes are saved into project/workspace
            if in_workspace_mode():
                save_workspace_from_session()
            else:
                save_session_into_project()

            draft_txt = st.session_state.main_text or ""
            sb_dict = {
                "synopsis": st.session_state.synopsis,
                "genre_style_notes": st.session_state.genre_style_notes,
                "world": st.session_state.world,
                "characters": st.session_state.characters,
                "outline": st.session_state.outline,
            }
            
            # Word count
            word_count = len(draft_txt.split())
            
            st.caption(f"**Draft Stats:** {word_count:,} words • Bay: {st.session_state.active_bay}")
            
            # Professional formatting options
            if st.session_state.active_bay == "FINAL":
                st.info("✨ **FINAL Bay** - Professional manuscript formatting available")
            
            # Manuscript formatting inputs
            with st.expander("📖 Manuscript Formatting Options", expanded=st.session_state.active_bay == "FINAL"):
                author_name = st.text_input("Author Name", value="Author Name", key="export_author")
                st.caption("Professional manuscript format follows industry standards (Shunn format)")

            st.subheader("Standard Exports")
            col_exp1, col_exp2 = st.columns(2)
            
            with col_exp1:
                st.download_button("Download Draft (.txt)", data=draft_txt, file_name=f"{stem}_draft.txt", mime="text/plain", use_container_width=True)
            with col_exp2:
                st.download_button(
                    "Download Draft (.md)",
                    data=draft_markdown(title, draft_txt, meta),
                    file_name=f"{stem}_draft.md",
                    mime="text/markdown",
                    use_container_width=True
                )
            
            st.download_button(
                "Download Story Bible (.md)",
                data=story_bible_markdown(title, sb_dict, meta),
                file_name=f"{stem}_story_bible.md",
                mime="text/markdown",
                use_container_width=True
            )
            
            st.subheader("📖 Professional Formats")
            
            # Manuscript format
            manuscript_txt = format_manuscript_standard(title, st.session_state.get("export_author", "Author Name"), draft_txt, word_count)
            col_man1, col_man2 = st.columns(2)
            with col_man1:
                st.download_button(
                    "📝 Manuscript (.txt)",
                    data=manuscript_txt,
                    file_name=f"{stem}_manuscript.txt",
                    mime="text/plain",
                    help="Industry-standard manuscript format (Shunn format)",
                    use_container_width=True
                )
            
            # eBook HTML
            ebook_html = format_ebook_html(title, st.session_state.get("export_author", "Author Name"), draft_txt)
            with col_man2:
                st.download_button(
                    "📱 eBook (.html)",
                    data=ebook_html,
                    file_name=f"{stem}_ebook.html",
                    mime="text/html",
                    help="Clean HTML for ebook conversion (import into Calibre/Kindle Create)",
                    use_container_width=True
                )
            
            st.caption("💡 **Tip:** Import .html into Calibre or Kindle Create to generate EPUB/MOBI files")

            st.divider()
            st.subheader("� Audio Export (Text-to-Speech)")
            
            st.caption("📢 Generate audio narration of your work for accessibility or review")
            
            # TTS Settings
            tts_col1, tts_col2 = st.columns(2)
            with tts_col1:
                tts_voice = st.selectbox(
                    "Voice",
                    ["Default", "Male", "Female"],
                    help="Browser-based text-to-speech voice",
                    key="tts_voice_select"
                )
            with tts_col2:
                tts_rate = st.slider(
                    "Speed",
                    0.5, 2.0, 1.0, 0.1,
                    help="Reading speed (1.0 = normal)",
                    key="tts_rate_select"
                )
            
            # Browser TTS using JavaScript
            tts_text = draft_txt[:5000] if len(draft_txt) > 5000 else draft_txt  # Limit for demo
            if st.button("🎧 Read Aloud (Preview)", key="tts_preview", use_container_width=True):
                # JavaScript to trigger browser TTS
                voice_setting = "default" if tts_voice == "Default" else "male" if tts_voice == "Male" else "female"
                
                st.markdown(
                    f"""
                    <script>
                    (function() {{
                        const text = {json.dumps(tts_text)};
                        const rate = {tts_rate};
                        
                        if ('speechSynthesis' in window) {{
                            // Cancel any ongoing speech
                            window.speechSynthesis.cancel();
                            
                            const utterance = new SpeechSynthesisUtterance(text);
                            utterance.rate = rate;
                            
                            // Try to select appropriate voice
                            const voices = window.speechSynthesis.getVoices();
                            if (voices.length > 0) {{
                                if ("{voice_setting}" === "male") {{
                                    const maleVoice = voices.find(v => v.name.toLowerCase().includes('male') || v.name.includes('David') || v.name.includes('James'));
                                    if (maleVoice) utterance.voice = maleVoice;
                                }} else if ("{voice_setting}" === "female") {{
                                    const femaleVoice = voices.find(v => v.name.toLowerCase().includes('female') || v.name.includes('Samantha') || v.name.includes('Karen'));
                                    if (femaleVoice) utterance.voice = femaleVoice;
                                }}
                            }}
                            
                            window.speechSynthesis.speak(utterance);
                        }} else {{
                            alert('Text-to-speech not supported in this browser');
                        }}
                    }})();
                    </script>
                    """,
                    unsafe_allow_html=True
                )
                st.success(f"🎧 Reading first {len(tts_text.split())} words aloud...")
            
            # Stop button
            if st.button("⏹️ Stop Audio", key="tts_stop", use_container_width=True):
                st.markdown(
                    """
                    <script>
                    if ('speechSynthesis' in window) {
                        window.speechSynthesis.cancel();
                    }
                    </script>
                    """,
                    unsafe_allow_html=True
                )
            
            st.caption("💡 **Tip:** For full audio export, use online TTS services like [Natural Reader](https://www.naturalreaders.com/) or [TTSReader](https://ttsreader.com/) to convert text to MP3")

            st.divider()
            st.subheader("�📄 DOCX Export")
            
            try:
                from docx import Document  # type: ignore
                from docx.shared import Pt, Inches  # type: ignore
                from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

                def _docx_bytes(doc: "Document") -> bytes:
                    buf = BytesIO()
                    doc.save(buf)
                    return buf.getvalue()

                # Standard draft DOCX
                d = Document()
                d.add_heading(f"Draft — {title}", level=0)
                for mk, mv in meta.items():
                    d.add_paragraph(f"{mk}: {mv}")
                d.add_paragraph(now_ts())
                d.add_paragraph("")
                for para in _split_paragraphs(draft_txt):
                    d.add_paragraph(para)
                
                col_doc1, col_doc2 = st.columns(2)
                with col_doc1:
                    st.download_button(
                        "Download Draft (.docx)",
                        data=_docx_bytes(d),
                        file_name=f"{stem}_draft.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                
                # Professional manuscript DOCX
                md = Document()
                # Set up manuscript formatting
                section = md.sections[0]
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                
                # Title page
                p = md.add_paragraph()
                p.add_run(st.session_state.get("export_author", "Author Name")).bold = True
                p.add_run(f"\n{word_count:,} words")
                
                md.add_page_break()
                
                # Title
                title_para = md.add_heading(title.upper(), level=1)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                by_para = md.add_paragraph(f"by {st.session_state.get('export_author', 'Author Name')}")
                by_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                md.add_page_break()
                
                # Body
                for para_text in _split_paragraphs(draft_txt):
                    if para_text.isupper() or para_text.startswith("Chapter") or para_text.startswith("CHAPTER"):
                        # Chapter heading
                        ch = md.add_heading(para_text, level=2)
                        ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        # Regular paragraph
                        p = md.add_paragraph(para_text)
                        p_format = p.paragraph_format
                        p_format.first_line_indent = Inches(0.5)
                        p_format.space_after = Pt(0)
                
                with col_doc2:
                    st.download_button(
                        "📝 Manuscript (.docx)",
                        data=_docx_bytes(md),
                        file_name=f"{stem}_manuscript.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        help="Professional manuscript format with proper margins and spacing",
                        use_container_width=True
                    )
            except Exception as e:
                st.caption("DOCX export unavailable (python-docx not installed).")

        with tab_bundle:
            st.caption("Bundle exports are merge-safe (imports never wipe your library).")
            if st.session_state.project_id:
                pj = json.dumps(make_project_bundle(st.session_state.project_id), ensure_ascii=False, indent=2)
                st.download_button(
                    "Download Project Bundle (.json)",
                    data=pj,
                    file_name=f"{_safe_filename(st.session_state.project_title,'project')}_bundle.json",
                    mime="application/json",
                )
            lib = json.dumps(make_library_bundle(), ensure_ascii=False, indent=2)
            st.download_button("Download Full Library (.json)", data=lib, file_name="olivetti_library_bundle.json", mime="application/json")

            st.divider()
            target_bay = st.selectbox("Imported projects go to bay", BAYS, index=0, key="io_bundle_target")
            rename = st.text_input("Optional rename (single project import)", key="io_bundle_rename")
            upb = st.file_uploader("Upload .json bundle", type=["json"], key="io_bundle_upload")
            switch_after = st.checkbox("Switch to imported project after import", value=True, key="io_bundle_switch")

            if st.button("Import Bundle", key="io_bundle_import"):
                if upb is None:
                    st.session_state.tool_output = "Import bundle: upload a .json file first."
                    autosave()
                else:
                    raw = upb.getvalue()
                    if raw is not None and len(raw) > MAX_UPLOAD_BYTES:
                        st.session_state.tool_output = "Import bundle: file too large."
                        autosave()
                    else:
                        try:
                            obj = json.loads((raw or b"").decode("utf-8"))
                        except Exception:
                            obj = None
                        if isinstance(obj, dict) and obj.get("projects"):
                            n = import_library_bundle(obj)
                            st.session_state.voice_status = f"Imported library bundle: {n} projects merged."
                            st.session_state.last_action = "Import Library Bundle"
                            autosave()
                        elif isinstance(obj, dict) and obj.get("project"):
                            pid = import_project_bundle(obj, target_bay=rename, rename=rename)
                            if pid:
                                st.session_state.voice_status = f"Imported project bundle → {pid}"
                                st.session_state.last_action = "Import Project Bundle"
                                autosave()
                                if switch_after:
                                    switch_bay(target_bay)
                                    load_project_into_session(pid)
                                    st.session_state.voice_status = f"{target_bay}: {st.session_state.project_title} (imported)"
                                    autosave()
                                    st.rerun()
                            else:
                                st.session_state.tool_output = "Import bundle: JSON did not look like a project bundle."
                                autosave()
                        else:
                            st.session_state.tool_output = "Import bundle: bundle type not recognized."
                            autosave()

    # Junk Drawer + Story Bible sections (labels hidden safely)
    with st.expander("🗃 Junk Drawer"):
        st.text_area(
            "Junk Drawer",
            key="junk",
            height=80,
            on_change=autosave,
            label_visibility="collapsed",
            help="Commands: /create: Title  |  /promote  |  /find: term",
        )
        st.text_area("Tool Output", value=st.session_state.tool_output, height=140, disabled=True)

    with st.expander("📝 Synopsis"):
        st.text_area("Synopsis", key="synopsis", height=100, on_change=autosave, label_visibility="collapsed", disabled=sb_locked)
        if not sb_locked and has_openai_key():
            if st.button("✨ Generate Synopsis", key="gen_synopsis"):
                generate_story_bible_section("Synopsis")
                st.rerun()

    with st.expander("🎭 Genre / Style Notes"):
        st.text_area(
            "Genre / Style Notes",
            key="genre_style_notes",
            height=80,
            on_change=autosave,
            label_visibility="collapsed",
            disabled=sb_locked,
        )
        if not sb_locked and has_openai_key():
            if st.button("✨ Generate Genre/Style", key="gen_genre"):
                generate_story_bible_section("Genre/Style")
                st.rerun()

    with st.expander("🌍 World Elements"):
        st.text_area("World", key="world", height=100, on_change=autosave, label_visibility="collapsed", disabled=sb_locked)
        if not sb_locked and has_openai_key():
            if st.button("✨ Generate World", key="gen_world"):
                generate_story_bible_section("World")
                st.rerun()

    with st.expander("👤 Characters"):
        st.text_area(
            "Characters",
            key="characters",
            height=120,
            on_change=autosave,
            label_visibility="collapsed",
            disabled=sb_locked,
        )
        if not sb_locked and has_openai_key():
            if st.button("✨ Generate Characters", key="gen_characters"):
                generate_story_bible_section("Characters")
                st.rerun()

    with st.expander("🧱 Outline"):
        st.text_area("Outline", key="outline", height=160, on_change=autosave, label_visibility="collapsed", disabled=sb_locked)
        if not sb_locked and has_openai_key():
            if st.button("✨ Generate Outline", key="gen_outline"):
                generate_story_bible_section("Outline")
                st.rerun()


# ============================================================
# CENTER — WRITING DESK
# ============================================================
with center:
    st.subheader("✍️ Writing Desk")
    
    # Voice Bible Status - Show what's controlling AI generation
    vb_status = _get_voice_bible_summary()
    st.caption(f"🎚️ **Active Controls:** {vb_status}")
    
    # Canon Guardian + Heatmap row
    col_tools1, col_tools2, col_tools3, col_tools4 = st.columns([2, 1, 2, 1])
    with col_tools1:
        st.checkbox("📖 Canon Guardian", key="canon_guardian_on", help="Real-time continuity validation against Story Bible")
    with col_tools2:
        if st.button("🔍 Check", key="btn_check_canon", disabled=not st.session_state.canon_guardian_on, use_container_width=True):
            text = (st.session_state.main_text or "").strip()
            if text:
                st.session_state.canon_issues = analyze_canon_conformity(text)
                error_count = sum(1 for i in st.session_state.canon_issues if i['severity'] == 'error')
                warn_count = sum(1 for i in st.session_state.canon_issues if i['severity'] == 'warning')
                st.session_state.ui_notice = f"📖 Found {error_count} error(s), {warn_count} warning(s)"
            else:
                st.session_state.ui_notice = "⚠️ No text to check"
                st.session_state.canon_issues = []
    
    with col_tools3:
        st.checkbox("🔥 Voice Heatmap", key="show_voice_heatmap", help="Analyze how well your draft follows Voice Bible controls")
    with col_tools4:
        if st.button("🔄 Analyze", key="btn_analyze_heatmap", disabled=not st.session_state.show_voice_heatmap, use_container_width=True):
            text = (st.session_state.main_text or "").strip()
            if text:
                st.session_state.voice_heatmap_data = analyze_voice_conformity(text)
                st.session_state.ui_notice = f"✅ Analyzed {len(st.session_state.voice_heatmap_data)} paragraph(s)"
            else:
                st.session_state.ui_notice = "⚠️ No text to analyze"
                st.session_state.voice_heatmap_data = []
    
    # Show color-coded text view when heatmap is enabled
    if st.session_state.show_voice_heatmap and st.session_state.voice_heatmap_data:
        st.caption("📊 **Live Heatmap View** (Green = On target, Yellow = Minor issues, Red = Major deviation)")
        
        # Build colored HTML view
        html_parts = ['<div style="background-color: #1e1e1e; padding: 15px; border-radius: 5px; font-family: monospace; line-height: 1.8; max-height: 650px; overflow-y: auto;">']
        
        for para_data in st.session_state.voice_heatmap_data:
            score = para_data["score"]
            # Determine color based on score
            if score >= 85:
                bg_color = "rgba(40, 167, 69, 0.3)"  # Green with transparency
                border_color = "#28a745"
            elif score >= 70:
                bg_color = "rgba(255, 193, 7, 0.3)"  # Yellow with transparency
                border_color = "#ffc107"
            elif score >= 50:
                bg_color = "rgba(253, 126, 20, 0.3)"  # Orange with transparency
                border_color = "#fd7e14"
            else:
                bg_color = "rgba(220, 53, 69, 0.3)"  # Red with transparency
                border_color = "#dc3545"
            
            # Add paragraph with inline background color
            issues_text = f" • {', '.join(para_data['issues'])}" if para_data['issues'] else ""
            html_parts.append(
                f'<div style="background-color: {bg_color}; border-left: 3px solid {border_color}; padding: 8px 12px; margin: 4px 0; border-radius: 3px; color: #e0e0e0;">'
                f'<span style="font-size: 0.8em; color: {border_color}; font-weight: bold;">[{score:.0f}/100{issues_text}]</span><br>'
                f'{para_data["text"]}'
                f'</div>'
            )
        
        html_parts.append('</div>')
        st.markdown(''.join(html_parts), unsafe_allow_html=True)
        
        st.caption("⚠️ Edit in text area below, then click 'Analyze' to update heatmap")
    
    # Display Canon Guardian issues if enabled
    if st.session_state.canon_guardian_on and st.session_state.canon_issues:
        error_count = sum(1 for i in st.session_state.canon_issues if i['severity'] == 'error')
        warn_count = sum(1 for i in st.session_state.canon_issues if i['severity'] == 'warning')
        
        st.warning(f"📖 **Canon Guardian**: {error_count} error(s), {warn_count} warning(s) detected")
        
        for idx, issue in enumerate(st.session_state.canon_issues[:5]):  # Show top 5
            severity_icon = "🔴" if issue['severity'] == 'error' else "🟡"
            confidence_color = "#dc3545" if issue['severity'] == 'error' else "#ffc107"
            
            with st.expander(f"{severity_icon} {issue['type'].replace('_', ' ').title()} (Confidence: {issue['confidence']}%)", expanded=idx==0):
                st.markdown(
                    f'<div style="background-color: rgba(220, 53, 69, 0.1); padding: 10px; border-radius: 4px; border-left: 3px solid {confidence_color};">'
                    f'<strong>Issue:</strong> {issue["issue"]}<br><br>'
                    f'<strong>Text:</strong> "{issue["text_snippet"]}..."'
                    f'</div>',
                    unsafe_allow_html=True
                )
                
                # Resolution buttons
                st.caption("**Resolution Options:**")
                res_cols = st.columns(len(issue['resolution_options']))
                for col_idx, option in enumerate(issue['resolution_options']):
                    if res_cols[col_idx].button(option, key=f"resolve_{idx}_{col_idx}", use_container_width=True):
                        if option == "Ignore":
                            # Add to ignored flags
                            flag_id = f"{issue['type']}_{issue['paragraph_index']}_{issue['issue'][:30]}"
                            if "canon_ignored_flags" not in st.session_state:
                                st.session_state.canon_ignored_flags = []
                            st.session_state.canon_ignored_flags.append(flag_id)
                            st.session_state.ui_notice = f"✅ Ignored: {issue['type']}"
                            st.rerun()
                        elif option == "Update Story Bible":
                            st.session_state.ui_notice = "💡 Navigate to Story Bible section to update canon"
                        elif option == "Fix Draft":
                            st.session_state.ui_notice = "💡 Edit the highlighted paragraph in your draft"
                        elif option == "Update Outline":
                            st.session_state.ui_notice = "💡 Navigate to Story Bible → Outline to update"
                        elif option == "Update World":
                            st.session_state.ui_notice = "💡 Navigate to Story Bible → World to update"
        
        if len(st.session_state.canon_issues) > 5:
            st.caption(f"... and {len(st.session_state.canon_issues) - 5} more issues")
    
    # Text area for editing (always shown)

    # --- Deferred main_text update to avoid Streamlit session state conflict ---
    if st.session_state.get("_defer_main_text_update", False):
        st.session_state.main_text = st.session_state._defer_main_text_value
        st.session_state._defer_main_text_update = False
        st.session_state._defer_main_text_value = ""
        autosave()
        st.rerun()


    st.text_area("Draft (main writing area)", key="main_text", height=400, on_change=autosave, label_visibility="collapsed")

    # SELECTION INPUT (paste text you want to modify)
    st.markdown("---")
    st.caption("**Selection** — Paste text here to work on it specifically (leave empty to work on full draft):")
    if "selection_text" not in st.session_state:
        st.session_state.selection_text = ""
    st.text_area("Selection (for targeted editing)", key="selection_text", height=100, placeholder="Paste the text you want to rewrite, expand, or modify...", label_visibility="collapsed")


    # --- AI PREVIEW WINDOW (shows AI output before accepting) ---
    if "ai_preview" not in st.session_state:
        st.session_state.ai_preview = ""
    if "ai_preview_action" not in st.session_state:
        st.session_state.ai_preview_action = ""

    # Show AI preview in the main writing desk panel
    if st.session_state.ai_preview:
        st.markdown("---")
        st.info(f"**Original Text** — for comparison:")
        orig_text = st.session_state.get('selection_text') or st.session_state.get('main_text') or ''
        st.markdown(
            f'<div style="background-color: #fffbe6; padding: 12px; border-radius: 5px; max-height: 200px; overflow-y: auto; white-space: pre-wrap; font-family: monospace; border: 1px solid #ffe58f; margin-bottom: 18px;">{orig_text}</div>',
            unsafe_allow_html=True
        )
        st.success(f"✨ **AI Generated Output** ({st.session_state.ai_preview_action}) — Review and accept or reject below:")
        st.markdown(
            f'<div style="background-color: #f0f2f6; padding: 15px; border-radius: 5px; max-height: 400px; overflow-y: auto; white-space: pre-wrap; font-family: monospace;">{st.session_state.ai_preview}</div>',
            unsafe_allow_html=True
        )

        accept_cols = st.columns([1, 1, 3])

        if accept_cols[0].button("✅ Accept", key="accept_ai", use_container_width=True):
            # Defer main_text update to avoid session state conflict
            if st.session_state.ai_preview_action in ["Write"]:
                # Append
                if (st.session_state.main_text or "").strip():
                    new_text = (st.session_state.main_text.rstrip() + "\n\n" + st.session_state.ai_preview).strip()
                else:
                    new_text = st.session_state.ai_preview
            else:
                # For selection-based edits, just show in preview - user copies manually
                # For full draft edits, replace
                if not st.session_state.selection_text.strip():
                    new_text = st.session_state.ai_preview
                else:
                    new_text = st.session_state.main_text

            st.session_state._defer_main_text_update = True
            st.session_state._defer_main_text_value = new_text
            st.session_state.voice_status = f"{st.session_state.ai_preview_action} accepted ✓"
            st.session_state.ai_preview = ""
            st.session_state.ai_preview_action = ""
            # Note: selection_text cannot be cleared here as it's bound to a widget
            # User can manually clear it if needed
            st.rerun()

        if accept_cols[1].button("❌ Reject", key="reject_ai", use_container_width=True):
            st.session_state.voice_status = f"{st.session_state.ai_preview_action} rejected"
            st.session_state.ai_preview = ""
            st.session_state.ai_preview_action = ""
            st.rerun()

        if st.session_state.selection_text.strip():
            accept_cols[2].info("💡 Copy the preview text and paste it back into your draft where you want it.")

    # Primary Actions (all respect Voice Bible + AI Intensity)
    b1 = st.columns(5)
    if b1[0].button("Write", key="btn_write", help="Continue writing (Voice Bible controlled)"):
        queue_action("Write")
    if b1[1].button("Rewrite", key="btn_rewrite", help="Rewrite for quality (Voice Bible controlled)"):
        queue_action("Rewrite")
    if b1[2].button("Expand", key="btn_expand", help="Add depth (Voice Bible controlled)"):
        queue_action("Expand")
    if b1[3].button("Rephrase", key="btn_rephrase", help="Rephrase last sentence (Voice Bible controlled)"):
        queue_action("Rephrase")
    if b1[4].button("Describe", key="btn_describe", help="Add description (Voice Bible controlled)"):
        queue_action("Describe")

    # Secondary Actions
    b2 = st.columns(5)
    if b2[0].button("Spell", key="btn_spell", help="Fix spelling/grammar (Voice Bible controlled)"):
        queue_action("Spell")
        st.rerun()
    if b2[1].button("Grammar", key="btn_grammar", help="Copyedit grammar (Voice Bible controlled)"):
        queue_action("Grammar")
        st.rerun()
    if b2[2].button("Find", key="btn_find", help="Search tool (see Junk Drawer)"):
        queue_action("__FIND_HINT__")
        st.rerun()
    if b2[3].button("Synonym", key="btn_synonym", help="Get synonyms for last word (Voice Bible aware)"):
        queue_action("Synonym")
        st.rerun()
    if b2[4].button("Sentence", key="btn_sentence", help="Rewrite last sentence options (Voice Bible aware)"):
        queue_action("Sentence")
        st.rerun()


# ============================================================
# RIGHT — VOICE BIBLE (AUTHOR ENGINE)
# ============================================================
with right:
    st.subheader("🎙 Voice Bible")
    st.caption("Your intelligent, trainable author engine. Each control adapts to your writing.")
    
    # Voice Bible Status Summary
    active_controls = []
    if st.session_state.vb_style_on:
        active_controls.append(f"Style: {st.session_state.writing_style}")
    if st.session_state.vb_genre_on:
        active_controls.append(f"Genre: {st.session_state.genre}")
    if st.session_state.vb_trained_on and st.session_state.trained_voice != "— None —":
        active_controls.append(f"Voice: {st.session_state.trained_voice}")
    if st.session_state.vb_match_on:
        active_controls.append("Match: ON")
    if st.session_state.vb_lock_on:
        active_controls.append("🔒 Lock: ON")
    
    if active_controls:
        st.info(f"**Active:** {' • '.join(active_controls)}")
    else:
        st.warning("⚠️ No Voice Bible controls active. AI will use defaults.")
    
    st.divider()

    with st.expander("✍️ Writing Style Engine", expanded=True):
        st.checkbox("Enable Writing Style", key="vb_style_on", on_change=autosave)
        style_col1, style_col2 = st.columns([2, 1])
        with style_col1:
            st.selectbox(
                "Writing Style (trainable)",
                ENGINE_STYLES,
                key="writing_style",
                disabled=not st.session_state.vb_style_on,
                on_change=autosave,
                help="All ENGINE styles are trainable with your own samples."
            )
        with style_col2:
            current_intensity = st.session_state.style_intensity if st.session_state.vb_style_on else 0.0
            intensity_label = "Subtle" if current_intensity < 0.35 else "Balanced" if current_intensity < 0.7 else "Strong"
            st.metric("Style", intensity_label, f"{current_intensity:.0%}")
        st.slider(
            "Style Intensity (distinctiveness)",
            0.0, 1.0,
            key="style_intensity",
            disabled=not st.session_state.vb_style_on,
            on_change=autosave,
            help="Lower = conservative, Higher = distinctive style"
        )

    with st.expander("🎨 Style Trainer (Adaptive Learning)", expanded=False):
        st.caption("🧠 Train ENGINE styles with your own writing. The system adapts to your voice patterns.")
        s_cols = st.columns([1.2, 1.0, 1.0])
        st_style = s_cols[0].selectbox(
            "Engine style (to train)",
            ENGINE_STYLES,
            key="style_train_style",
            help="Choose which ENGINE style to train with your samples"
        )
        st_lane = s_cols[1].selectbox(
            "Lane (writing mode)",
            LANES,
            key="style_train_lane",
            help="Train for specific writing modes: Dialogue, Narration, Interiority, or Action"
        )
        split_mode = s_cols[2].selectbox(
            "Split (sample granularity)",
            ["Paragraphs", "Whole"],
            key="style_train_split",
            help="Paragraphs = multiple samples, Whole = single sample"
        )
        bank = (st.session_state.get("style_banks") or {}).get(st_style, {})
        lanes = bank.get("lanes") or {}
        counts = {ln: len((lanes.get(ln) or [])) for ln in LANES}
        total_samples = sum(counts.values())
        if total_samples > 0:
            st.success(f"✅ {st_style} has {total_samples} training samples across all lanes")
        else:
            st.warning(f"⚠️ {st_style} has no training samples yet. Add samples to activate this style.")
        up = st.file_uploader(
            "Upload training (.txt/.md/.docx)",
            type=["txt", "md", "docx"],
            key="style_train_upload",
            help="Upload a document containing your writing style examples"
        )
        paste = st.text_area(
            "Paste training text (style)",
            key="style_train_paste",
            height=140,
            placeholder="Paste 1-3 paragraphs of your best writing in this style...",
            help="AI uses these samples to learn your style. Be consistent and clear."
        )
        c1, c2, c3 = st.columns([1, 1, 1])
        if c1.button("Add Samples", key="style_train_add", use_container_width=True):
            ftxt, fname = _read_uploaded_text(up)
            src = _normalize_text((paste or "").strip() if (paste or "").strip() else ftxt)
            if not src.strip():
                st.session_state.tool_output = "Style Trainer: no text provided (or file too large)."
                st.session_state.voice_status = "Style Trainer blocked"
                autosave()
            else:
                n = add_style_samples(st_style, st_lane, src, split_mode=split_mode)
                st.session_state.voice_status = f"Style Trainer: added {n} sample(s) → {st_style} • {st_lane}"
                st.session_state.tool_output = _clamp_text(f"✅ Added {n} sample(s) to {st_style} / {st_lane}.\n\nThe engine will now adapt to these patterns when generating {st_lane} content.\n\nSource: {fname or 'paste'}")
                autosave()
                st.rerun()
        if c2.button("Delete last", key="style_train_del", use_container_width=True):
            if delete_last_style_sample(st_style, st_lane):
                st.session_state.voice_status = f"Style Trainer: deleted last → {st_style} • {st_lane}"
                autosave()
                st.rerun()
            else:
                st.warning("Nothing to delete for that style/lane.")

        if c3.button("Clear text", key="style_train_clear", use_container_width=True):
            st.session_state.ui_notice = "Style trainer text cleared."
            queue_action("__STYLE_CLEAR_PASTE__")
            st.rerun()

        # Enhanced lane statistics with visual feedback
        st.caption("📊 Training samples per lane:")
        cols_display = st.columns(4)
        for idx, ln in enumerate(LANES):
            with cols_display[idx]:
                count = counts[ln]
                status = "✅" if count >= 3 else "⚠️" if count > 0 else "➖"
                st.metric(ln, count, status)
        if st.button("Clear THIS lane", key="style_train_clear_lane", use_container_width=True):
            clear_style_lane(st_style, st_lane)
            st.session_state.voice_status = f"Style Trainer: cleared lane → {st_style} • {st_lane}"
            autosave()
            st.rerun()


    st.divider()
    with st.expander("🎭 Genre Intelligence", expanded=False):
        st.caption("🎯 AI adapts genre markers, pacing, and tone automatically.")
        col_g1, col_g2 = st.columns([2, 1])
        with col_g1:
            st.checkbox("Enable Genre Influence", key="vb_genre_on", on_change=autosave)
        with col_g2:
            genre_strength = st.session_state.genre_intensity if st.session_state.vb_genre_on else 0.0
            st.metric("Genre", f"{genre_strength:.0%}")
        st.selectbox(
            "Genre (affects pacing/tone)",
            ["Literary", "Noir", "Thriller", "Comedy", "Lyrical", "Horror", "Romance", "SciFi", "Fantasy"],
            key="genre",
            disabled=not st.session_state.vb_genre_on,
            on_change=autosave,
            help="Influences pacing, vocabulary, and tonal decisions"
        )
        st.slider(
            "Genre Intensity (conventions)",
            0.0, 1.0,
            key="genre_intensity",
            disabled=not st.session_state.vb_genre_on,
            on_change=autosave,
            help="How strongly genre conventions influence the writing"
        )

    st.divider()
    with st.expander("🧬 Trained Voice (Vector Matching)", expanded=False):
        st.caption("🎓 AI retrieves your best examples contextually using semantic search.")
        col_t1, col_t2 = st.columns([2, 1])
        with col_t1:
            st.checkbox("Enable Trained Voice", key="vb_trained_on", on_change=autosave)
        with col_t2:
            trained_str = st.session_state.trained_intensity if st.session_state.vb_trained_on else 0.0
            st.metric("Voice", f"{trained_str:.0%}")
        trained_options = voice_names_for_selector()
        if st.session_state.trained_voice not in trained_options:
            st.session_state.trained_voice = "— None —"
        st.selectbox(
            "Trained Voice (from Voice Vault)",
            trained_options,
            key="trained_voice",
            disabled=not st.session_state.vb_trained_on,
            on_change=autosave,
            help="Select a voice you've trained in the Voice Vault"
        )
        st.slider(
            "Trained Voice Intensity (mimicry)",
            0.0, 1.0,
            key="trained_intensity",
            disabled=not st.session_state.vb_trained_on,
            on_change=autosave,
            help="How closely AI mimics your trained voice patterns"
        )

    with st.expander("🧬 Voice Vault (Training Samples)", expanded=False):
        st.caption("🎯 Smart semantic retrieval: AI finds your most relevant examples automatically.")

        existing_voices = [v for v in (st.session_state.voices or {}).keys()]
        existing_voices = sorted(existing_voices, key=lambda x: (x not in ("Voice A", "Voice B"), x))
        if not existing_voices:
            existing_voices = ["Voice A", "Voice B"]


        vcol1, vcol2 = st.columns([2, 1])
        vault_voice = vcol1.selectbox(
            "Vault voice (choose to train/manage)",
            existing_voices,
            key="vault_voice_sel",
            help="Select a voice to train or manage"
        )
        new_name = vcol2.text_input(
            "New voice name",
            key="vault_new_voice",
            label_visibility="collapsed",
            placeholder="New voice name"
        )
        if vcol2.button("Create", key="vault_create_voice"):
            if create_custom_voice(new_name):
                st.session_state.voice_status = f"Voice created: {new_name.strip()}"
                autosave()
                st.rerun()
            else:
                st.warning("Could not create that voice (empty or already exists).")

        v = (st.session_state.voices or {}).get(vault_voice, {})
        lane_counts = {ln: len((v.get("lanes", {}) or {}).get(ln, []) or []) for ln in LANES}
        total_vault = sum(lane_counts.values())
        if total_vault > 0:
            st.success(f"✅ '{vault_voice}' has {total_vault} samples. AI will adapt to these patterns.")
        else:
            st.info(f"ℹ️ '{vault_voice}' has no samples yet. Add samples to train this voice.")

        lane = st.selectbox("Lane (writing mode)", LANES, key="vault_lane_sel", help="Choose which writing mode to train")
        sample = st.text_area(
            "Sample (best writing in this lane)",
            key="vault_sample_text",
            height=140,
            label_visibility="collapsed",
            placeholder="Paste a passage from your best writing in this lane...",
            help="Add examples of your writing. AI uses semantic matching to find relevant samples."
        )
        a1, a2 = st.columns([1, 1])
        if a1.button("Add sample", key="vault_add_sample", use_container_width=True):
            if add_voice_sample(vault_voice, lane, sample):
                st.session_state.ui_notice = f"✅ Added sample → {vault_voice} • {lane}"
                autosave()
                st.rerun()
            else:
                st.warning("No sample text found.")
        st.caption("📊 Sample distribution:")
        cols_v = st.columns(4)
        for idx, ln in enumerate(LANES):
            with cols_v[idx]:
                c = lane_counts[ln]
                status = "✅" if c >= 3 else "⚠️" if c > 0 else "➖"
                st.metric(ln, c, status)
        if a2.button("Delete last sample", key="vault_del_last", use_container_width=True):
            if delete_voice_sample(vault_voice, lane, index_from_end=0):
                st.session_state.voice_status = f"Deleted last sample → {vault_voice} • {lane}"
                autosave()
                st.rerun()
            else:
                st.warning("Nothing to delete for that lane.")

    st.divider()

    # ============ MATCH MY STYLE (One-Shot Adaptation) ============
    with st.expander("✨ Match My Style (One-Shot)", expanded=False):
        st.caption("🎨 AI adapts immediately to a single style example. Paste up to 2000 words for deep analysis.")
        
        col_m1, col_m2 = st.columns([2, 1])
        with col_m1:
            st.checkbox("Enable Match My Style", key="vb_match_on", on_change=autosave)
        with col_m2:
            match_str = st.session_state.match_intensity if st.session_state.vb_match_on else 0.0
            st.metric("Match", f"{match_str:.0%}")
        
        st.text_area(
            "Style Example (up to 2000 words)",
            key="voice_sample",
            height=300,
            max_chars=15000,
            disabled=not st.session_state.vb_match_on,
            on_change=autosave,
            placeholder="Paste your best writing here (up to 2000 words). Click 'Analyze' to identify the strongest passages...",
            help="AI will adapt its output to match this example's patterns. Analysis highlights your strongest sentences."
        )
        
        col_a1, col_a2 = st.columns([1, 1])
        with col_a1:
            if st.button("🔍 Analyze Style", key="analyze_style_btn", use_container_width=True, disabled=not st.session_state.vb_match_on):
                text = (st.session_state.voice_sample or "").strip()
                if text:
                    samples = analyze_style_samples(text)
                    st.session_state.analyzed_style_samples = samples
                    if samples:
                        st.session_state.ui_notice = f"✅ Found {len(samples)} strong writing samples"
                    else:
                        st.session_state.ui_notice = "⚠️ Not enough text to analyze. Add more sentences."
                else:
                    st.session_state.ui_notice = "⚠️ No text to analyze"
                    st.session_state.analyzed_style_samples = []
        
        # Show analysis results
        if st.session_state.analyzed_style_samples:
            st.success(f"📊 **Top {len(st.session_state.analyzed_style_samples)} Strongest Writing Samples** (by score)")
            for idx, sample in enumerate(st.session_state.analyzed_style_samples[:5], 1):
                with st.container():
                    st.markdown(f"**#{idx}** (Score: {sample['score']:.1f})")
                    st.info(f"✨ {sample['text']}")
                    st.caption(f"📈 {sample['words']} words • Vocab: {sample['unique_ratio']:.0%} • Sensory: {sample['sensory']} • Verbs: {sample['verbs']}")
        
        st.slider(
            "Match Intensity", 
            0.0, 1.0, 
            key="match_intensity", 
            disabled=not st.session_state.vb_match_on, 
            on_change=autosave,
            help="How closely AI mimics the provided example"
        )


    st.divider()
    with st.expander("🔒 Voice Lock (Hard Constraint)", expanded=False):
        st.caption("⚠️ MANDATORY directives. Use for absolute rules like 'never use passive voice'.")
        col_l1, col_l2 = st.columns([2, 1])
        with col_l1:
            st.checkbox("Voice Lock (Hard Constraint)", key="vb_lock_on", on_change=autosave)
        with col_l2:
            lock_str = st.session_state.lock_intensity if st.session_state.vb_lock_on else 0.0
            st.metric("Lock", f"{lock_str:.0%}")
        st.text_area(
            "Voice Lock Prompt (absolute constraints)",
            key="voice_lock_prompt",
            height=80,
            disabled=not st.session_state.vb_lock_on,
            on_change=autosave,
            placeholder="Example: 'Never use adverbs. Always use short sentences. No passive voice.'",
            help="Absolute constraints. AI will enforce these rules strictly."
        )
        st.slider(
            "Lock Strength", 0.0, 1.0, key="lock_intensity", disabled=not st.session_state.vb_lock_on, on_change=autosave, help="Higher = stricter enforcement of constraints"
        )

    st.divider()
    with st.expander("⚙️ Technical Controls (POV/Tense)", expanded=False):
        st.caption("🎯 Enforces point of view and verb tense consistency across all AI generation.")
        col_tc1, col_tc2 = st.columns([2, 1])
        with col_tc1:
            st.checkbox("Enable Technical Controls", key="vb_technical_on", on_change=autosave)
        with col_tc2:
            tech_str = st.session_state.technical_intensity if st.session_state.vb_technical_on else 0.0
            st.metric("Tech", f"{tech_str:.0%}")
        col_tech1, col_tech2 = st.columns(2)
        with col_tech1:
            st.selectbox(
                "POV (point of view)",
                ["First", "Close Third", "Omniscient"],
                key="pov",
                disabled=not st.session_state.vb_technical_on,
                on_change=autosave,
                help="Point of view for narrative perspective"
            )
        with col_tech2:
            st.selectbox(
                "Tense (verb tense)",
                ["Past", "Present"],
                key="tense",
                disabled=not st.session_state.vb_technical_on,
                on_change=autosave,
                help="Verb tense for all actions"
            )
        st.slider(
            "Technical Enforcement",
            0.0, 1.0,
            key="technical_intensity",
            disabled=not st.session_state.vb_technical_on,
            on_change=autosave,
            help="How strictly AI enforces POV/Tense rules. Higher = more rigid enforcement."
        )

# SAFETY NET SAVE EVERY RERUN
save_all_to_disk()
